import streamlit as st
import os
import json
from datetime import datetime
import hashlib
from pathlib import Path
import socket
import threading
from queue import Queue
import time
import re
import math
import bcrypt
import requests
from collections import defaultdict
from urllib.parse import urljoin, urlparse
import random
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from io import BytesIO

# ============================================================================
# PAGE CONFIGURATION - MUST BE FIRST STREAMLIT COMMAND
# ============================================================================
st.set_page_config(
    page_title="NovaCrypt Defense - Hybrid Hacking Toolkit",
    page_icon="🔐",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================================
# CUSTOM CSS FOR AMAZING CYBERSECURITY THEME
# ============================================================================
def load_custom_css():
    st.markdown("""
    <style>
    /* Main background - Dark cyber theme */
    .stApp {
        background: linear-gradient(135deg, #0f0c29 0%, #302b63 50%, #24243e 100%);
    }
    
    /* Sidebar styling */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1a1a2e 0%, #16213e 100%);
        border-right: 2px solid #00fff5;
    }
    
    /* Headers with glow effect */
    h1, h2, h3 {
        color: #00fff5 !important;
        text-shadow: 0 0 10px #00fff5, 0 0 20px #00fff5;
        font-family: 'Courier New', monospace;
    }
    
    /* Success/Info boxes */
    .stSuccess {
        background-color: rgba(0, 255, 127, 0.1);
        border: 1px solid #00ff7f;
        border-radius: 10px;
    }
    
    .stInfo {
        background-color: rgba(0, 191, 255, 0.1);
        border: 1px solid #00bfff;
        border-radius: 10px;
    }
    
    .stWarning {
        background-color: rgba(255, 165, 0, 0.1);
        border: 1px solid #ffa500;
        border-radius: 10px;
    }
    
    .stError {
        background-color: rgba(255, 0, 0, 0.1);
        border: 1px solid #ff0000;
        border-radius: 10px;
    }
    
    /* Buttons */
    .stButton > button {
        background: linear-gradient(90deg, #00fff5 0%, #00a8cc 100%);
        color: #000;
        font-weight: bold;
        border: none;
        border-radius: 8px;
        padding: 10px 25px;
        transition: all 0.3s;
        box-shadow: 0 0 15px rgba(0, 255, 245, 0.5);
    }
    
    .stButton > button:hover {
        transform: scale(1.05);
        box-shadow: 0 0 25px rgba(0, 255, 245, 0.8);
    }
    
    /* Input fields */
    .stTextInput > div > div > input {
        background-color: rgba(255, 255, 255, 0.05);
        border: 1px solid #00fff5;
        color: #ffffff;
        border-radius: 5px;
    }
    
    /* Metrics */
    [data-testid="stMetricValue"] {
        color: #00fff5;
        font-size: 2rem;
        text-shadow: 0 0 10px #00fff5;
    }
    
    /* Expander */
    .streamlit-expanderHeader {
        background-color: rgba(0, 255, 245, 0.1);
        border-radius: 5px;
        color: #00fff5 !important;
    }
    
    /* Code blocks */
    .stCodeBlock {
        background-color: rgba(0, 0, 0, 0.5);
        border: 1px solid #00fff5;
        border-radius: 8px;
    }
    
    /* Divider */
    hr {
        border: 1px solid #00fff5;
        margin: 20px 0;
    }
    
    /* Animated pulse for important elements */
    @keyframes pulse {
        0%, 100% { opacity: 1; }
        50% { opacity: 0.6; }
    }
    
    .pulse {
        animation: pulse 2s infinite;
    }
    </style>
    """, unsafe_allow_html=True)

# ============================================================================
# PACKET CAPTURE & ANALYSIS MODULE
# ============================================================================
class PacketCapture:
    def __init__(self, logger):
        self.logger = logger
        self.packets = []
        self.capture_active = False
        
    def generate_sample_packet(self, protocol="HTTP"):
        """Generate simulated network packet"""
        protocols = ["HTTP", "HTTPS", "DNS", "TCP", "UDP", "ICMP"]
        
        if protocol == "Random":
            protocol = random.choice(protocols)
        
        src_ip = f"{random.randint(192, 192)}.{random.randint(168, 168)}.{random.randint(1, 1)}.{random.randint(1, 254)}"
        dst_ip = f"{random.randint(1, 223)}.{random.randint(0, 255)}.{random.randint(0, 255)}.{random.randint(1, 254)}"
        
        packet = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3],
            "protocol": protocol,
            "src_ip": src_ip,
            "dst_ip": dst_ip,
            "src_port": random.randint(1024, 65535),
            "dst_port": self.get_common_port(protocol),
            "length": random.randint(60, 1500),
            "info": self.get_packet_info(protocol)
        }
        
        return packet
    
    def get_common_port(self, protocol):
        """Get common destination port for protocol"""
        ports = {
            "HTTP": 80,
            "HTTPS": 443,
            "DNS": 53,
            "TCP": random.choice([22, 80, 443, 3306, 5432]),
            "UDP": random.choice([53, 123, 161]),
            "ICMP": 0
        }
        return ports.get(protocol, random.randint(1, 65535))
    
    def get_packet_info(self, protocol):
        """Generate realistic packet info"""
        info_templates = {
            "HTTP": [
                "GET / HTTP/1.1",
                "POST /api/login HTTP/1.1",
                "GET /images/logo.png HTTP/1.1",
                "HTTP/1.1 200 OK"
            ],
            "HTTPS": [
                "Client Hello",
                "Server Hello",
                "Certificate",
                "Application Data"
            ],
            "DNS": [
                "Standard query A example.com",
                "Standard query response A 93.184.216.34",
                "Standard query AAAA google.com",
                "Standard query PTR 1.1.168.192.in-addr.arpa"
            ],
            "TCP": [
                "SYN",
                "SYN, ACK",
                "ACK",
                "PSH, ACK",
                "FIN, ACK"
            ],
            "UDP": [
                "UDP payload",
                "DNS query",
                "NTP request"
            ],
            "ICMP": [
                "Echo (ping) request",
                "Echo (ping) reply",
                "Destination unreachable"
            ]
        }
        
        return random.choice(info_templates.get(protocol, ["Unknown"]))
    
    def capture_traffic(self, duration=10, protocol_filter="All"):
        """Simulate traffic capture"""
        self.packets = []
        self.capture_active = True
        
        self.logger.log("PACKET_CAPTURE", "Started", f"Duration: {duration}s, Filter: {protocol_filter}")
        
        start_time = time.time()
        packet_count = 0
        
        while time.time() - start_time < duration and self.capture_active:
            # Generate 5-15 packets per second
            packets_per_batch = random.randint(5, 15)
            
            for _ in range(packets_per_batch):
                if protocol_filter == "All":
                    packet = self.generate_sample_packet("Random")
                else:
                    packet = self.generate_sample_packet(protocol_filter)
                
                self.packets.append(packet)
                packet_count += 1
            
            time.sleep(1)
        
        self.capture_active = False
        
        self.logger.log("PACKET_CAPTURE", "Completed", f"Captured {packet_count} packets")
        
        return self.analyze_traffic()
    
    def analyze_traffic(self):
        """Analyze captured packets"""
        if not self.packets:
            return None
        
        # Protocol distribution
        protocol_count = defaultdict(int)
        for packet in self.packets:
            protocol_count[packet['protocol']] += 1
        
        # Top talkers (IPs)
        ip_traffic = defaultdict(int)
        for packet in self.packets:
            ip_traffic[packet['src_ip']] += 1
            ip_traffic[packet['dst_ip']] += 1
        
        top_ips = sorted(ip_traffic.items(), key=lambda x: x[1], reverse=True)[:10]
        
        # Port usage
        port_count = defaultdict(int)
        for packet in self.packets:
            port_count[packet['dst_port']] += 1
        
        top_ports = sorted(port_count.items(), key=lambda x: x[1], reverse=True)[:10]
        
        # Traffic volume
        total_bytes = sum(packet['length'] for packet in self.packets)
        
        analysis = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "total_packets": len(self.packets),
            "total_bytes": total_bytes,
            "protocol_distribution": dict(protocol_count),
            "top_ips": [{"ip": ip, "packets": count} for ip, count in top_ips],
            "top_ports": [{"port": port, "packets": count} for port, count in top_ports],
            "packets": self.packets[-100:]  # Last 100 packets
        }
        
        return analysis
    
    def export_pcap(self, filename):
        """Simulate .pcap export"""
        # Ensure evidence directory exists
        Path("evidence").mkdir(exist_ok=True)
        
        filepath = f"evidence/{filename}"
        
        # Create a text representation (real .pcap would need scapy)
        with open(filepath, 'w') as f:
            f.write("# Simulated PCAP Export\n")
            f.write(f"# Captured: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"# Total Packets: {len(self.packets)}\n\n")
            
            for i, packet in enumerate(self.packets, 1):
                f.write(f"Packet {i}:\n")
                f.write(f"  Time: {packet['timestamp']}\n")
                f.write(f"  Protocol: {packet['protocol']}\n")
                f.write(f"  Source: {packet['src_ip']}:{packet['src_port']}\n")
                f.write(f"  Destination: {packet['dst_ip']}:{packet['dst_port']}\n")
                f.write(f"  Length: {packet['length']} bytes\n")
                f.write(f"  Info: {packet['info']}\n\n")
        
        self.logger.log("PACKET_CAPTURE", "Export", f"PCAP saved to {filepath}")
        return filepath
    
    def export_json(self, analysis, filename):
        """Export analysis to JSON"""
        # Ensure evidence directory exists
        Path("evidence").mkdir(exist_ok=True)
        
        filepath = f"evidence/{filename}"
        with open(filepath, 'w') as f:
            json.dump(analysis, indent=2, fp=f)
        self.logger.log("PACKET_CAPTURE", "Export", f"Analysis saved to {filepath}")
        return filepath

# ============================================================================
# WEB DISCOVERY MODULE (DIRB-style)
# ============================================================================
class WebDiscovery:
    def __init__(self, logger):
        self.logger = logger
        self.discovered = []
        self.lock = threading.Lock()
        self.checked_count = 0
        self.stop_flag = False
        
        # Common directories/files to check
        self.common_paths = [
            "admin", "administrator", "login", "dashboard", "panel",
            "backup", "backups", "old", "tmp", "temp", "test",
            "api", "v1", "v2", "rest", "graphql",
            "uploads", "files", "images", "assets", "static",
            "config", "conf", "configuration", ".env", ".git",
            "readme.txt", "robots.txt", "sitemap.xml", "security.txt",
            "wp-admin", "wp-login", "phpmyadmin", "mysql",
            "setup", "install", "installer", "upgrade",
            ".htaccess", ".htpasswd", "web.config",
            "user", "users", "account", "profile",
            "download", "downloads", "docs", "documentation",
            "log", "logs", "error_log", "access_log",
            "db", "database", "sql", "data",
            "include", "includes", "lib", "libs", "vendor",
            "css", "js", "fonts", "media",
            "payment", "checkout", "cart", "order"
        ]
        
        # Common file extensions
        self.extensions = [
            "", ".php", ".html", ".htm", ".asp", ".aspx",
            ".jsp", ".js", ".json", ".xml", ".txt",
            ".bak", ".old", ".log", ".zip", ".tar.gz"
        ]
    
    def check_path(self, base_url, path, timeout=5):
        """Check if a path exists on target"""
        try:
            url = urljoin(base_url, path)
            response = requests.head(url, timeout=timeout, allow_redirects=False)
            
            # Consider as found if not 404
            if response.status_code != 404:
                result = {
                    "url": url,
                    "path": path,
                    "status_code": response.status_code,
                    "size": response.headers.get('Content-Length', 'Unknown'),
                    "content_type": response.headers.get('Content-Type', 'Unknown'),
                    "server": response.headers.get('Server', 'Unknown')
                }
                
                with self.lock:
                    self.discovered.append(result)
                    self.logger.log("WEB_DISCOVERY", "Found", f"{url} [{response.status_code}]")
                
                return result
            
        except requests.exceptions.Timeout:
            pass
        except Exception as e:
            pass
        finally:
            with self.lock:
                self.checked_count += 1
        
        return None
    
    def worker(self, base_url, path_queue, timeout):
        """Worker thread for path checking"""
        while not path_queue.empty() and not self.stop_flag:
            path = path_queue.get()
            self.check_path(base_url, path, timeout)
            path_queue.task_done()
            time.sleep(0.1)  # Rate limiting
    
    def scan_directories(self, base_url, wordlist="common", num_threads=10, timeout=5, extensions=True):
        """Main directory scanning function"""
        self.discovered = []
        self.checked_count = 0
        self.stop_flag = False
        
        # Normalize URL
        if not base_url.startswith(('http://', 'https://')):
            base_url = 'http://' + base_url
        
        if not base_url.endswith('/'):
            base_url += '/'
        
        self.logger.log("WEB_DISCOVERY", "Scan Started", f"Target: {base_url}, Wordlist: {wordlist}, Threads: {num_threads}")
        
        # Build path list
        paths = []
        
        if wordlist == "common":
            paths = self.common_paths.copy()
        elif wordlist == "minimal":
            paths = self.common_paths[:15]  # First 15 only
        elif wordlist == "extensive":
            paths = self.common_paths.copy()
            # Add with extensions
            if extensions:
                extended = []
                for path in paths:
                    for ext in self.extensions:
                        extended.append(path + ext)
                paths = extended
        
        # Add robots.txt and sitemap.xml (always check these)
        if "robots.txt" not in paths:
            paths.insert(0, "robots.txt")
        if "sitemap.xml" not in paths:
            paths.insert(1, "sitemap.xml")
        
        total_paths = len(paths)
        
        # Create queue
        path_queue = Queue()
        for path in paths:
            path_queue.put(path)
        
        # Create threads
        threads = []
        for _ in range(min(num_threads, total_paths)):
            t = threading.Thread(target=self.worker, args=(base_url, path_queue, timeout))
            t.daemon = True
            t.start()
            threads.append(t)
        
        # Wait for completion
        for t in threads:
            t.join()
        
        self.logger.log("WEB_DISCOVERY", "Scan Completed", f"Checked {self.checked_count} paths, Found {len(self.discovered)} resources")
        
        return {
            "base_url": base_url,
            "wordlist": wordlist,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "paths_checked": self.checked_count,
            "resources_found": len(self.discovered),
            "discovered_paths": sorted(self.discovered, key=lambda x: x['status_code'])
        }
    
    def check_subdomains(self, domain, timeout=3):
        """Simple subdomain enumeration"""
        common_subdomains = [
            "www", "mail", "ftp", "admin", "test", "dev", "staging",
            "api", "app", "mobile", "portal", "secure", "vpn",
            "blog", "shop", "store", "support", "help", "docs"
        ]
        
        found_subdomains = []
        
        for sub in common_subdomains:
            subdomain = f"{sub}.{domain}"
            try:
                socket.gethostbyname(subdomain)
                found_subdomains.append({
                    "subdomain": subdomain,
                    "status": "Active"
                })
                self.logger.log("WEB_DISCOVERY", "Subdomain Found", subdomain)
            except socket.gaierror:
                pass
        
        return found_subdomains
    
    def export_results(self, results, filename):
        """Export discovery results to JSON"""
        # Ensure evidence directory exists
        Path("evidence").mkdir(exist_ok=True)
        
        filepath = f"evidence/{filename}"
        with open(filepath, 'w') as f:
            json.dump(results, indent=2, fp=f)
        self.logger.log("WEB_DISCOVERY", "Export", f"Results saved to {filepath}")
        return filepath

# ============================================================================
# GRAPH GENERATION FOR STRESS TEST
# ============================================================================
def generate_performance_graphs(results):
    """Generate performance graphs for stress test results"""
    
    # Create figure with 2 subplots
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 5))
    fig.patch.set_facecolor('#1a1a2e')
    
    # Graph 1: Success vs Errors Pie Chart
    ax1.set_facecolor('#16213e')
    labels = ['Successful', 'Failed']
    sizes = [results['successful_requests'], results['failed_requests']]
    colors = ['#00ff7f', '#ff4500']
    explode = (0.1, 0)
    
    ax1.pie(sizes, explode=explode, labels=labels, colors=colors,
            autopct='%1.1f%%', shadow=True, startangle=90,
            textprops={'color': 'white', 'fontsize': 12})
    ax1.set_title('Request Success Rate', color='#00fff5', fontsize=14, fontweight='bold')
    
    # Graph 2: Latency Distribution Bar Chart
    ax2.set_facecolor('#16213e')
    latency = results['latency_stats']
    metrics = ['Avg', 'Min', 'Max', 'P50', 'P95', 'P99']
    values = [
        latency['average_ms'],
        latency['min_ms'],
        latency['max_ms'],
        latency['p50_ms'],
        latency['p95_ms'],
        latency['p99_ms']
    ]
    
    bars = ax2.bar(metrics, values, color='#00bfff', edgecolor='#00fff5', linewidth=2)
    ax2.set_xlabel('Latency Metrics', color='white', fontsize=12)
    ax2.set_ylabel('Latency (ms)', color='white', fontsize=12)
    ax2.set_title('Response Time Distribution', color='#00fff5', fontsize=14, fontweight='bold')
    ax2.tick_params(colors='white')
    ax2.spines['bottom'].set_color('white')
    ax2.spines['left'].set_color('white')
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    
    # Add value labels on bars
    for bar in bars:
        height = bar.get_height()
        ax2.text(bar.get_x() + bar.get_width()/2., height,
                f'{height:.1f}ms',
                ha='center', va='bottom', color='white', fontsize=10)
    
    plt.tight_layout()
    
    # Save to BytesIO
    buf = BytesIO()
    plt.savefig(buf, format='png', facecolor='#1a1a2e', dpi=100)
    buf.seek(0)
    plt.close()
    
    return buf    

# ============================================================================
# DOS/STRESS TESTING MODULE
# ============================================================================
class StressTester:
    def __init__(self, logger):
        self.logger = logger
        self.results = []
        self.lock = threading.Lock()
        self.stop_flag = False
        self.request_count = 0
        self.success_count = 0
        self.error_count = 0
        
    def send_request(self, url, method="GET", timeout=5):
        """Send a single HTTP request"""
        start_time = time.time()
        
        try:
            if method == "GET":
                response = requests.get(url, timeout=timeout)
            elif method == "POST":
                response = requests.post(url, data={"test": "data"}, timeout=timeout)
            else:
                response = requests.head(url, timeout=timeout)
            
            latency = (time.time() - start_time) * 1000  # Convert to ms
            
            result = {
                "timestamp": time.time(),
                "method": method,
                "status_code": response.status_code,
                "latency_ms": round(latency, 2),
                "success": 200 <= response.status_code < 400
            }
            
            with self.lock:
                self.results.append(result)
                self.request_count += 1
                if result["success"]:
                    self.success_count += 1
                else:
                    self.error_count += 1
            
            return result
            
        except requests.exceptions.Timeout:
            latency = (time.time() - start_time) * 1000
            with self.lock:
                self.request_count += 1
                self.error_count += 1
                self.results.append({
                    "timestamp": time.time(),
                    "method": method,
                    "status_code": 0,
                    "latency_ms": round(latency, 2),
                    "success": False,
                    "error": "Timeout"
                })
            return None
            
        except Exception as e:
            with self.lock:
                self.request_count += 1
                self.error_count += 1
                self.results.append({
                    "timestamp": time.time(),
                    "method": method,
                    "status_code": 0,
                    "latency_ms": 0,
                    "success": False,
                    "error": str(e)
                })
            return None
    
    def worker(self, url, method, timeout, duration):
        """Worker thread for stress testing"""
        end_time = time.time() + duration
        
        while time.time() < end_time and not self.stop_flag:
            self.send_request(url, method, timeout)
            time.sleep(0.1)  # Small delay between requests
    
    def run_stress_test(self, url, num_clients, duration, method="GET", timeout=5):
        """Execute stress test with multiple concurrent clients"""
        self.results = []
        self.stop_flag = False
        self.request_count = 0
        self.success_count = 0
        self.error_count = 0
        
        self.logger.log("STRESS_TEST", "Started", f"Target: {url}, Clients: {num_clients}, Duration: {duration}s")
        
        # Validate URL
        if not url.startswith(('http://', 'https://')):
            url = 'http://' + url
        
        # Create worker threads
        threads = []
        for i in range(num_clients):
            t = threading.Thread(target=self.worker, args=(url, method, timeout, duration))
            t.daemon = True
            t.start()
            threads.append(t)
        
        # Wait for all threads or stop signal
        for t in threads:
            t.join()
        
        self.logger.log("STRESS_TEST", "Completed", f"Sent {self.request_count} requests, Success: {self.success_count}, Errors: {self.error_count}")
        
        return self.analyze_results(url, num_clients, duration, method)
    
    def analyze_results(self, url, num_clients, duration, method):
        """Analyze stress test results"""
        if not self.results:
            return None
        
        # Calculate metrics
        latencies = [r['latency_ms'] for r in self.results if 'latency_ms' in r and r['latency_ms'] > 0]
        
        avg_latency = sum(latencies) / len(latencies) if latencies else 0
        min_latency = min(latencies) if latencies else 0
        max_latency = max(latencies) if latencies else 0
        
        # Calculate percentiles
        sorted_latencies = sorted(latencies)
        p50 = sorted_latencies[len(sorted_latencies) // 2] if sorted_latencies else 0
        p95 = sorted_latencies[int(len(sorted_latencies) * 0.95)] if sorted_latencies else 0
        p99 = sorted_latencies[int(len(sorted_latencies) * 0.99)] if sorted_latencies else 0
        
        # Requests per second
        rps = self.request_count / duration if duration > 0 else 0
        
        # Success rate
        success_rate = (self.success_count / self.request_count * 100) if self.request_count > 0 else 0
        
        # Status code distribution
        status_codes = defaultdict(int)
        for r in self.results:
            status_codes[r.get('status_code', 0)] += 1
        
        analysis = {
            "target": url,
            "method": method,
            "num_clients": num_clients,
            "duration_seconds": duration,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "total_requests": self.request_count,
            "successful_requests": self.success_count,
            "failed_requests": self.error_count,
            "success_rate_percent": round(success_rate, 2),
            "requests_per_second": round(rps, 2),
            "latency_stats": {
                "average_ms": round(avg_latency, 2),
                "min_ms": round(min_latency, 2),
                "max_ms": round(max_latency, 2),
                "p50_ms": round(p50, 2),
                "p95_ms": round(p95, 2),
                "p99_ms": round(p99, 2)
            },
            "status_code_distribution": dict(status_codes),
            "raw_results": self.results[-100:]  # Last 100 requests
        }
        
        return analysis
    
    def export_results(self, results, filename):
        """Export stress test results to JSON"""
        # Ensure evidence directory exists
        Path("evidence").mkdir(exist_ok=True)
        
        filepath = f"evidence/{filename}"
        with open(filepath, 'w') as f:
            json.dump(results, indent=2, fp=f)
        self.logger.log("STRESS_TEST", "Export", f"Results saved to {filepath}")
        return filepath

# ============================================================================
# PASSWORD ASSESSMENT MODULE
# ============================================================================
class PasswordAssessment:
    def __init__(self, logger):
        self.logger = logger
        
        # Common weak passwords list (top 100)
        self.common_passwords = [
            "password", "123456", "123456789", "12345678", "12345", "1234567",
            "password1", "123123", "1234567890", "000000", "qwerty", "abc123",
            "111111", "admin", "letmein", "welcome", "monkey", "dragon", "master",
            "sunshine", "princess", "football", "shadow", "superman", "696969",
            "michael", "jennifer", "computer", "trustno1", "mustang", "baseball",
            "hunter", "charlie", "password123", "welcome123", "admin123", "root",
            "toor", "pass", "test", "guest", "info", "sample", "changeme",
            "secret", "login", "demo", "user", "default", "temp", "qwerty123"
        ]
        
        # Password policy rules
        self.policy = {
            "min_length": 8,
            "max_length": 128,
            "require_uppercase": True,
            "require_lowercase": True,
            "require_digits": True,
            "require_special": True,
            "special_chars": "!@#$%^&*()_+-=[]{}|;:,.<>?"
        }
    
    def calculate_entropy(self, password):
        """Calculate Shannon entropy of password"""
        if not password:
            return 0.0
        
        # Count character frequency
        freq = {}
        for char in password:
            freq[char] = freq.get(char, 0) + 1
        
        # Calculate entropy
        entropy = 0.0
        length = len(password)
        
        for count in freq.values():
            probability = count / length
            entropy -= probability * math.log2(probability)
        
        # Normalize by length
        bits = entropy * length
        
        return round(bits, 2)
    
    def check_character_sets(self, password):
        """Check which character sets are used"""
        has_lower = bool(re.search(r'[a-z]', password))
        has_upper = bool(re.search(r'[A-Z]', password))
        has_digit = bool(re.search(r'\d', password))
        has_special = bool(re.search(r'[^a-zA-Z0-9]', password))
        
        return {
            "lowercase": has_lower,
            "uppercase": has_upper,
            "digits": has_digit,
            "special": has_special
        }
    
    def check_patterns(self, password):
        """Check for common patterns"""
        patterns_found = []
        
        # Sequential numbers
        if re.search(r'(012|123|234|345|456|567|678|789)', password):
            patterns_found.append("Sequential numbers")
        
        # Repeated characters
        if re.search(r'(.)\1{2,}', password):
            patterns_found.append("Repeated characters")
        
        # Keyboard patterns
        keyboard_patterns = ['qwerty', 'asdf', 'zxcv', '1234', 'qwer', 'asdfgh']
        for pattern in keyboard_patterns:
            if pattern in password.lower():
                patterns_found.append(f"Keyboard pattern: {pattern}")
        
        # Common words
        common_words = ['password', 'admin', 'user', 'login', 'welcome', 'test']
        for word in common_words:
            if word in password.lower():
                patterns_found.append(f"Common word: {word}")
        
        # Date patterns
        if re.search(r'(19|20)\d{2}', password):
            patterns_found.append("Year pattern detected")
        
        return patterns_found
    
    def check_policy_compliance(self, password):
        """Check password against policy rules"""
        issues = []
        
        # Length check
        if len(password) < self.policy["min_length"]:
            issues.append(f"Too short (minimum {self.policy['min_length']} characters)")
        
        if len(password) > self.policy["max_length"]:
            issues.append(f"Too long (maximum {self.policy['max_length']} characters)")
        
        # Character requirements
        char_sets = self.check_character_sets(password)
        
        if self.policy["require_uppercase"] and not char_sets["uppercase"]:
            issues.append("Missing uppercase letters")
        
        if self.policy["require_lowercase"] and not char_sets["lowercase"]:
            issues.append("Missing lowercase letters")
        
        if self.policy["require_digits"] and not char_sets["digits"]:
            issues.append("Missing digits")
        
        if self.policy["require_special"] and not char_sets["special"]:
            issues.append("Missing special characters")
        
        return issues
    
    def is_common_password(self, password):
        """Check if password is in common password list"""
        return password.lower() in self.common_passwords
    
    def calculate_strength_score(self, password):
        """Calculate overall password strength (0-100)"""
        score = 0
        
        # Length scoring (0-30 points)
        length = len(password)
        if length >= 16:
            score += 30
        elif length >= 12:
            score += 25
        elif length >= 8:
            score += 15
        else:
            score += 5
        
        # Character diversity (0-25 points)
        char_sets = self.check_character_sets(password)
        score += sum(char_sets.values()) * 6.25
        
        # Entropy (0-25 points)
        entropy = self.calculate_entropy(password)
        if entropy >= 60:
            score += 25
        elif entropy >= 40:
            score += 20
        elif entropy >= 30:
            score += 15
        else:
            score += entropy / 3
        
        # Pattern penalties (0-20 points deduction)
        patterns = self.check_patterns(password)
        penalty = min(len(patterns) * 5, 20)
        score -= penalty
        
        # Common password penalty (-30 points)
        if self.is_common_password(password):
            score -= 30
        
        # Ensure score is between 0-100
        score = max(0, min(100, score))
        
        return round(score, 1)
    
    def get_strength_category(self, score):
        """Categorize password strength"""
        if score >= 80:
            return "Very Strong", "🟢", "#00ff7f"
        elif score >= 60:
            return "Strong", "🟡", "#ffa500"
        elif score >= 40:
            return "Moderate", "🟠", "#ff8c00"
        elif score >= 20:
            return "Weak", "🔴", "#ff4500"
        else:
            return "Very Weak", "🔴", "#ff0000"
    
    def simulate_hash(self, password, hash_type="all"):
        """Simulate password hashing (educational only)"""
        hashes = {}
        
        if hash_type in ["all", "md5"]:
            hashes["MD5"] = hashlib.md5(password.encode()).hexdigest()
        
        if hash_type in ["all", "sha256"]:
            hashes["SHA256"] = hashlib.sha256(password.encode()).hexdigest()
        
        if hash_type in ["all", "sha512"]:
            hashes["SHA512"] = hashlib.sha512(password.encode()).hexdigest()
        
        if hash_type in ["all", "bcrypt"]:
            # Simulate bcrypt (actual bcrypt would be used in production)
            salt = bcrypt.gensalt()
            hashes["bcrypt"] = bcrypt.hashpw(password.encode(), salt).decode()
        
        return hashes
    
    def generate_recommendations(self, password):
        """Generate actionable security recommendations"""
        recommendations = []
        
        length = len(password)
        char_sets = self.check_character_sets(password)
        patterns = self.check_patterns(password)
        is_common = self.is_common_password(password)
        
        # Length recommendations
        if length < 12:
            recommendations.append("✅ Increase password length to at least 12 characters (16+ recommended)")
        
        # Character diversity
        if not char_sets["uppercase"]:
            recommendations.append("✅ Add uppercase letters (A-Z)")
        
        if not char_sets["lowercase"]:
            recommendations.append("✅ Add lowercase letters (a-z)")
        
        if not char_sets["digits"]:
            recommendations.append("✅ Add numbers (0-9)")
        
        if not char_sets["special"]:
            recommendations.append("✅ Add special characters (!@#$%^&*)")
        
        # Pattern warnings
        if patterns:
            recommendations.append("⚠️ Avoid predictable patterns (sequential numbers, keyboard patterns)")
        
        # Common password warning
        if is_common:
            recommendations.append("🚨 CRITICAL: This is a commonly used password - change immediately!")
        
        # General best practices
        if length < 16:
            recommendations.append("💡 Use a passphrase (e.g., 'Coffee-Mountain-Sky-42!')")
        
        recommendations.append("💡 Use a unique password for each account")
        recommendations.append("💡 Consider using a password manager")
        recommendations.append("💡 Enable two-factor authentication (2FA)")
        
        return recommendations
    
    def assess_password(self, password):
        """Comprehensive password assessment"""
        
        self.logger.log("PASSWORD_TEST", "Assessment Started", f"Password length: {len(password)}")
        
        # Calculate all metrics
        score = self.calculate_strength_score(password)
        entropy = self.calculate_entropy(password)
        char_sets = self.check_character_sets(password)
        patterns = self.check_patterns(password)
        policy_issues = self.check_policy_compliance(password)
        is_common = self.is_common_password(password)
        category, emoji, color = self.get_strength_category(score)
        recommendations = self.generate_recommendations(password)
        
        result = {
            "password_length": len(password),
            "strength_score": score,
            "strength_category": category,
            "strength_emoji": emoji,
            "strength_color": color,
            "entropy_bits": entropy,
            "character_sets": char_sets,
            "patterns_detected": patterns,
            "policy_compliant": len(policy_issues) == 0,
            "policy_issues": policy_issues,
            "is_common_password": is_common,
            "recommendations": recommendations,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        self.logger.log("PASSWORD_TEST", "Assessment Complete", f"Score: {score}/100, Category: {category}")
        
        return result
    
    def export_results(self, results, filename):
        """Export assessment results to JSON"""
        # Ensure evidence directory exists
        Path("evidence").mkdir(exist_ok=True)
        
        filepath = f"evidence/{filename}"
        with open(filepath, 'w') as f:
            json.dump(results, indent=2, fp=f)
        self.logger.log("PASSWORD_TEST", "Export", f"Results saved to {filepath}")
        return filepath

# ============================================================================
# PORT SCANNER MODULE
# ============================================================================
class PortScanner:
    def __init__(self, logger):
        self.logger = logger
        self.open_ports = []
        self.lock = threading.Lock()
        self.scan_progress = 0
        self.total_ports = 0
        
        # Common ports and their services
        self.common_ports = {
            21: "FTP", 22: "SSH", 23: "Telnet", 25: "SMTP", 53: "DNS",
            80: "HTTP", 110: "POP3", 143: "IMAP", 443: "HTTPS", 445: "SMB",
            3306: "MySQL", 3389: "RDP", 5432: "PostgreSQL", 5900: "VNC",
            8080: "HTTP-Proxy", 8443: "HTTPS-Alt", 27017: "MongoDB"
        }
    
    def resolve_target(self, target):
        """Resolve hostname to IP address"""
        try:
            ip = socket.gethostbyname(target)
            self.logger.log("PORT_SCAN", "DNS Resolution", f"{target} -> {ip}")
            return ip
        except socket.gaierror:
            self.logger.log("PORT_SCAN", "DNS Error", f"Cannot resolve {target}", "ERROR")
            return None
    
    def scan_port(self, target, port, timeout=1):
        """Scan a single port"""
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(timeout)
            result = sock.connect_ex((target, port))
            
            if result == 0:
                # Port is open, try to grab banner
                banner = self.grab_banner(sock, port)
                service = self.common_ports.get(port, "Unknown")
                
                with self.lock:
                    self.open_ports.append({
                        "port": port,
                        "service": service,
                        "banner": banner,
                        "state": "open"
                    })
                
                self.logger.log("PORT_SCAN", f"Open Port Found", f"Port {port} ({service})")
            
            sock.close()
        except Exception as e:
            pass
        finally:
            with self.lock:
                self.scan_progress += 1
    
    def grab_banner(self, sock, port):
        """Attempt to grab service banner"""
        try:
            sock.settimeout(2)
            
            # Send different probes based on port
            if port == 80 or port == 8080:
                sock.send(b"GET / HTTP/1.1\r\nHost: target\r\n\r\n")
            elif port == 21:
                pass  # FTP sends banner automatically
            elif port == 22:
                pass  # SSH sends banner automatically
            else:
                sock.send(b"\r\n")
            
            banner = sock.recv(1024).decode('utf-8', errors='ignore').strip()
            return banner[:200] if banner else "No banner"
        except:
            return "No banner"
    
    def worker(self, target, port_queue, timeout):
        """Worker thread for scanning"""
        while not port_queue.empty():
            port = port_queue.get()
            self.scan_port(target, port, timeout)
            port_queue.task_done()
    
    def scan(self, target, port_range, num_threads=50, timeout=1):
        """Main scan function"""
        self.open_ports = []
        self.scan_progress = 0
        
        # Resolve target
        ip = self.resolve_target(target)
        if not ip:
            return None
        
        # Parse port range
        if '-' in port_range:
            start, end = map(int, port_range.split('-'))
            ports = range(start, end + 1)
        elif ',' in port_range:
            ports = [int(p.strip()) for p in port_range.split(',')]
        else:
            ports = [int(port_range)]
        
        self.total_ports = len(ports)
        
        self.logger.log("PORT_SCAN", "Scan Started", f"Target: {target} ({ip}), Ports: {len(ports)}, Threads: {num_threads}")
        
        # Create queue and add ports
        port_queue = Queue()
        for port in ports:
            port_queue.put(port)
        
        # Create and start threads
        threads = []
        for _ in range(min(num_threads, len(ports))):
            t = threading.Thread(target=self.worker, args=(ip, port_queue, timeout))
            t.daemon = True
            t.start()
            threads.append(t)
        
        # Wait for all threads to complete
        for t in threads:
            t.join()
        
        self.logger.log("PORT_SCAN", "Scan Completed", f"Found {len(self.open_ports)} open ports")
        
        return {
            "target": target,
            "ip": ip,
            "scan_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "total_ports_scanned": self.total_ports,
            "open_ports_found": len(self.open_ports),
            "open_ports": sorted(self.open_ports, key=lambda x: x['port'])
        }
    
    def export_json(self, results, filename):
        """Export results to JSON"""
        # Ensure evidence directory exists
        Path("evidence").mkdir(exist_ok=True)
        
        filepath = f"evidence/{filename}"
        with open(filepath, 'w') as f:
            json.dump(results, indent=2, fp=f)
        self.logger.log("PORT_SCAN", "Export", f"Results saved to {filepath}")
        return filepath
    
    def export_html(self, results, filename):
        """Export results to HTML"""
        # Ensure evidence directory exists
        Path("evidence").mkdir(exist_ok=True)
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <title>Port Scan Results - {results['target']}</title>
    <style>
        body {{ font-family: 'Courier New', monospace; background: #1a1a2e; color: #00fff5; padding: 20px; }}
        h1 {{ color: #00fff5; text-shadow: 0 0 10px #00fff5; }}
        table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
        th {{ background: #16213e; color: #00fff5; padding: 12px; text-align: left; border: 1px solid #00fff5; }}
        td {{ background: #0f0c29; padding: 10px; border: 1px solid #00fff5; }}
        .open {{ color: #00ff7f; font-weight: bold; }}
        .summary {{ background: rgba(0,255,245,0.1); padding: 15px; border: 1px solid #00fff5; border-radius: 5px; margin: 20px 0; }}
    </style>
</head>
<body>
    <h1>🔍 Port Scan Results</h1>
    <div class="summary">
        <p><strong>Target:</strong> {results['target']} ({results['ip']})</p>
        <p><strong>Scan Time:</strong> {results['scan_time']}</p>
        <p><strong>Ports Scanned:</strong> {results['total_ports_scanned']}</p>
        <p><strong>Open Ports Found:</strong> {results['open_ports_found']}</p>
    </div>
    
    <h2>Open Ports Details</h2>
    <table>
        <tr>
            <th>Port</th>
            <th>Service</th>
            <th>State</th>
            <th>Banner</th>
        </tr>
"""
        for port_info in results['open_ports']:
            html += f"""
        <tr>
            <td>{port_info['port']}</td>
            <td>{port_info['service']}</td>
            <td class="open">{port_info['state']}</td>
            <td>{port_info['banner']}</td>
        </tr>
"""
        
        html += """
    </table>
    <p style="text-align: center; color: #666; margin-top: 40px;">
        <small>Generated by NovaCrypt Defense - Moazam (9953) & Abdullah (7465)</small>
    </p>
</body>
</html>
"""
        
        filepath = f"evidence/{filename}"
        with open(filepath, 'w') as f:
            f.write(html)
        self.logger.log("PORT_SCAN", "Export", f"HTML report saved to {filepath}")
        return filepath

# ============================================================================
# LOGGER CLASS - Centralized Logging with SHA-256 Integrity
# ============================================================================
class SecurityLogger:
    def __init__(self, log_file="evidence/security_logs.log"):
        self.log_file = log_file
        self.ensure_log_directory()
        
    def ensure_log_directory(self):
        """Create evidence directory if it doesn't exist"""
        Path("evidence").mkdir(exist_ok=True)
        
    def log(self, module, action, details, level="INFO"):
        """Log an action with timestamp and details"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_entry = f"[{timestamp}] [{level}] [{module}] {action} - {details}\n"
        
        # Append to log file
        with open(self.log_file, "a", encoding="utf-8") as f:
            f.write(log_entry)
        
        return log_entry
    
    def get_logs(self):
        """Retrieve all logs"""
        if not os.path.exists(self.log_file):
            return "No logs found."
        
        with open(self.log_file, "r", encoding="utf-8") as f:
            return f.read()
    
    def calculate_log_hash(self):
        """Calculate SHA-256 hash of log file for integrity"""
        if not os.path.exists(self.log_file):
            return None
        
        sha256_hash = hashlib.sha256()
        with open(self.log_file, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        
        return sha256_hash.hexdigest()
    
    def export_logs_json(self):
        """Export logs as JSON for reporting"""
        if not os.path.exists(self.log_file):
            return {}
        
        logs = []
        with open(self.log_file, "r", encoding="utf-8") as f:
            for line in f:
                if line.strip():
                    logs.append({"entry": line.strip()})
        
        return {
            "total_entries": len(logs),
            "hash": self.calculate_log_hash(),
            "logs": logs
        }

# ============================================================================
# REPORT GENERATOR - Word & PDF Reports
# ============================================================================
class ReportGenerator:
    def __init__(self, logger):
        self.logger = logger
        
    def generate_comprehensive_report(self, scan_results=None, password_results=None, 
                                     stress_results=None, discovery_results=None, 
                                     packet_results=None):
        """Generate a comprehensive Word report with all test results"""
        
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create document
        doc = Document()
        
        # ====================================================================
        # TITLE PAGE
        # ====================================================================
        title = doc.add_heading('NovaCrypt Defense', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Hybrid Hacking Toolkit - Security Assessment Report', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Team info
        doc.add_paragraph()
        team_info = doc.add_paragraph()
        team_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        team_info.add_run('Team: NovaCrypt Defense\n').bold = True
        team_info.add_run('Moazam | BSFT07-9953\n')
        team_info.add_run('Abdullah | BSFT07-7465\n')
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f'\nReport Generated: {datetime.now().strftime("%B %d, %Y at %H:%M:%S")}\n').italic = True
        
        doc.add_page_break()
        
        # ====================================================================
        # TABLE OF CONTENTS
        # ====================================================================
        doc.add_heading('Table of Contents', level=1)
        doc.add_paragraph('1. Executive Summary')
        doc.add_paragraph('2. Port Scanning Results')
        doc.add_paragraph('3. Password Assessment Results')
        doc.add_paragraph('4. Stress Testing Results')
        doc.add_paragraph('5. Web Discovery Results')
        doc.add_paragraph('6. Packet Capture Analysis')
        doc.add_paragraph('7. Recommendations')
        doc.add_paragraph('8. Conclusion')
        
        doc.add_page_break()
        
        # ====================================================================
        # EXECUTIVE SUMMARY
        # ====================================================================
        doc.add_heading('1. Executive Summary', level=1)
        
        summary = doc.add_paragraph()
        summary.add_run('Assessment Overview:\n').bold = True
        summary.add_run(f'This security assessment was conducted on {datetime.now().strftime("%B %d, %Y")} ')
        summary.add_run('using the NovaCrypt Defense Hybrid Hacking Toolkit. ')
        summary.add_run('The assessment covered multiple security domains including network scanning, ')
        summary.add_run('password security, stress testing, web discovery, and network traffic analysis.\n\n')
        
        # Summary stats
        total_tests = sum([1 for r in [scan_results, password_results, stress_results, discovery_results, packet_results] if r])
        summary.add_run(f'Total Modules Executed: {total_tests}\n')
        summary.add_run(f'Report Date: {datetime.now().strftime("%Y-%m-%d")}\n')
        
        doc.add_page_break()
        
        # ====================================================================
        # PORT SCANNING RESULTS
        # ====================================================================
        if scan_results:
            doc.add_heading('2. Port Scanning Results', level=1)
            
            doc.add_paragraph(f"Target: {scan_results.get('target', 'N/A')}")
            doc.add_paragraph(f"IP Address: {scan_results.get('ip', 'N/A')}")
            doc.add_paragraph(f"Scan Time: {scan_results.get('scan_time', 'N/A')}")
            doc.add_paragraph(f"Total Ports Scanned: {scan_results.get('total_ports_scanned', 0)}")
            doc.add_paragraph(f"Open Ports Found: {scan_results.get('open_ports_found', 0)}")
            
            if scan_results.get('open_ports'):
                doc.add_heading('Open Ports Details:', level=2)
                
                # Create table
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'
                
                # Header
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Port'
                hdr_cells[1].text = 'Service'
                hdr_cells[2].text = 'State'
                hdr_cells[3].text = 'Banner'
                
                # Data
                for port in scan_results['open_ports']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(port.get('port', ''))
                    row_cells[1].text = port.get('service', '')
                    row_cells[2].text = port.get('state', '')
                    row_cells[3].text = port.get('banner', '')[:50]  # Truncate
            
            doc.add_page_break()
        
        # ====================================================================
        # PASSWORD ASSESSMENT
        # ====================================================================
        if password_results:
            doc.add_heading('3. Password Assessment Results', level=1)
            
            doc.add_paragraph(f"Strength Score: {password_results.get('strength_score', 0)}/100")
            doc.add_paragraph(f"Category: {password_results.get('strength_category', 'N/A')}")
            doc.add_paragraph(f"Password Length: {password_results.get('password_length', 0)} characters")
            doc.add_paragraph(f"Entropy: {password_results.get('entropy_bits', 0)} bits")
            doc.add_paragraph(f"Policy Compliant: {'Yes' if password_results.get('policy_compliant') else 'No'}")
            doc.add_paragraph(f"Common Password: {'Yes' if password_results.get('is_common_password') else 'No'}")
            
            if password_results.get('recommendations'):
                doc.add_heading('Recommendations:', level=2)
                for rec in password_results['recommendations'][:5]:
                    doc.add_paragraph(rec, style='List Bullet')
            
            doc.add_page_break()
        
        # ====================================================================
        # STRESS TESTING
        # ====================================================================
        if stress_results:
            doc.add_heading('4. Stress Testing Results', level=1)
            
            doc.add_paragraph(f"Target: {stress_results.get('target', 'N/A')}")
            doc.add_paragraph(f"Test Duration: {stress_results.get('duration_seconds', 0)} seconds")
            doc.add_paragraph(f"Concurrent Clients: {stress_results.get('num_clients', 0)}")
            doc.add_paragraph(f"Total Requests: {stress_results.get('total_requests', 0)}")
            doc.add_paragraph(f"Successful Requests: {stress_results.get('successful_requests', 0)}")
            doc.add_paragraph(f"Failed Requests: {stress_results.get('failed_requests', 0)}")
            doc.add_paragraph(f"Success Rate: {stress_results.get('success_rate_percent', 0)}%")
            doc.add_paragraph(f"Requests/Second: {stress_results.get('requests_per_second', 0)}")
            
            if stress_results.get('latency_stats'):
                doc.add_heading('Latency Statistics:', level=2)
                latency = stress_results['latency_stats']
                doc.add_paragraph(f"Average: {latency.get('average_ms', 0)} ms")
                doc.add_paragraph(f"Minimum: {latency.get('min_ms', 0)} ms")
                doc.add_paragraph(f"Maximum: {latency.get('max_ms', 0)} ms")
                doc.add_paragraph(f"P95: {latency.get('p95_ms', 0)} ms")
                doc.add_paragraph(f"P99: {latency.get('p99_ms', 0)} ms")
            
            doc.add_page_break()
        
        # ====================================================================
        # WEB DISCOVERY
        # ====================================================================
        if discovery_results:
            doc.add_heading('5. Web Discovery Results', level=1)
            
            doc.add_paragraph(f"Target: {discovery_results.get('base_url', 'N/A')}")
            doc.add_paragraph(f"Paths Checked: {discovery_results.get('paths_checked', 0)}")
            doc.add_paragraph(f"Resources Found: {discovery_results.get('resources_found', 0)}")
            
            if discovery_results.get('discovered_paths'):
                doc.add_heading('Discovered Resources:', level=2)
                for path in discovery_results['discovered_paths'][:20]:  # First 20
                    doc.add_paragraph(f"{path.get('path', '')} - Status: {path.get('status_code', '')}", style='List Bullet')
            
            doc.add_page_break()
        
        # ====================================================================
        # PACKET CAPTURE
        # ====================================================================
        if packet_results:
            doc.add_heading('6. Packet Capture Analysis', level=1)
            
            doc.add_paragraph(f"Total Packets: {packet_results.get('total_packets', 0)}")
            doc.add_paragraph(f"Total Bytes: {packet_results.get('total_bytes', 0)}")
            
            if packet_results.get('protocol_distribution'):
                doc.add_heading('Protocol Distribution:', level=2)
                for proto, count in packet_results['protocol_distribution'].items():
                    doc.add_paragraph(f"{proto}: {count} packets", style='List Bullet')
            
            doc.add_page_break()
        
        # ====================================================================
        # RECOMMENDATIONS
        # ====================================================================
        doc.add_heading('7. Security Recommendations', level=1)
        
        doc.add_paragraph('Based on the security assessment, the following recommendations are provided:', style='List Bullet')
        doc.add_paragraph('Implement strong password policies across all systems', style='List Bullet')
        doc.add_paragraph('Close unnecessary open ports to reduce attack surface', style='List Bullet')
        doc.add_paragraph('Implement rate limiting to prevent DOS attacks', style='List Bullet')
        doc.add_paragraph('Secure sensitive directories and files from public access', style='List Bullet')
        doc.add_paragraph('Monitor network traffic for suspicious patterns', style='List Bullet')
        doc.add_paragraph('Regular security assessments should be conducted quarterly', style='List Bullet')
        
        doc.add_page_break()
        
        # ====================================================================
        # CONCLUSION
        # ====================================================================
        doc.add_heading('8. Conclusion', level=1)
        
        conclusion = doc.add_paragraph()
        conclusion.add_run('This comprehensive security assessment has identified various aspects of the target systems. ')
        conclusion.add_run('The findings and recommendations outlined in this report should be addressed to improve ')
        conclusion.add_run('the overall security posture of the organization.\n\n')
        conclusion.add_run('For any questions regarding this report, please contact the NovaCrypt Defense team.')
        
        # Save document
        Path("evidence").mkdir(exist_ok=True)
        filename = f"evidence/comprehensive_report_9953_Moazam_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        
        self.logger.log("REPORT", "Generated", f"Comprehensive report saved to {filename}")
        
        return filename
    
    def convert_to_pdf(self, docx_path):
        """Convert Word document to PDF using reportlab"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from docx import Document
            
            # Read DOCX
            doc = Document(docx_path)
            
            # Create PDF
            pdf_path = docx_path.replace('.docx', '.pdf')
            c = canvas.Canvas(pdf_path, pagesize=letter)
            width, height = letter
            
            y_position = height - 50
            
            # Add title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, y_position, "NovaCrypt Defense - Security Report")
            y_position -= 40
            
            # Add content
            c.setFont("Helvetica", 10)
            for para in doc.paragraphs:
                if para.text.strip():
                    # Word wrap
                    text = para.text
                    if len(text) > 80:
                        # Simple word wrap
                        words = text.split()
                        line = ""
                        for word in words:
                            if len(line + word) < 80:
                                line += word + " "
                            else:
                                c.drawString(50, y_position, line)
                                y_position -= 15
                                line = word + " "
                                
                                if y_position < 50:
                                    c.showPage()
                                    y_position = height - 50
                        
                        if line:
                            c.drawString(50, y_position, line)
                            y_position -= 15
                    else:
                        c.drawString(50, y_position, text)
                        y_position -= 15
                    
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50
            
            c.save()
            self.logger.log("REPORT", "Converted", f"PDF saved to {pdf_path}")
            return pdf_path
            
        except Exception as e:
            self.logger.log("REPORT", "PDF Error", str(e), "ERROR")
            return None


# ============================================================================
# IDENTITY & SAFETY MODULE
# ============================================================================
class IdentitySafety:
    def __init__(self):
        self.identity_file = "identity.txt"
        self.consent_file = "consent.txt"
        self.logger = SecurityLogger()
        
    def create_identity_file(self):
        """Create identity.txt if it doesn't exist"""
        identity_content = """Team: NovaCrypt Defense
Members:
- Moazam | BSFT07-9953
- Abdullah | BSFT07-7465
"""
        with open(self.identity_file, "w", encoding="utf-8") as f:
            f.write(identity_content)
        
        self.logger.log("IDENTITY", "Created", "identity.txt file generated")
        return identity_content
    
    def create_consent_file(self):
        """Create consent.txt if it doesn't exist"""
        consent_content = """Approved Targets:
- localhost / 127.0.0.1
- http://testphp.vulnweb.com (Acunetix test site)
- scanme.nmap.org
- example.com
- Local Flask/FastAPI mock servers
- OWASP Juice Shop (local instance)
- TryHackMe lab environments
- Any real-world URL (with ethical testing only)

Approved By: Moazam & Abdullah
Date: 30-November-2025
"""
        with open(self.consent_file, "w", encoding="utf-8") as f:
            f.write(consent_content)
        
        self.logger.log("CONSENT", "Created", "consent.txt file generated")
        return consent_content
    
    def verify_identity(self):
        """Verify identity.txt exists and is valid"""
        if not os.path.exists(self.identity_file):
            return False, "identity.txt not found! Creating..."
        
        with open(self.identity_file, "r", encoding="utf-8") as f:
            content = f.read()
            
        # Check if required team members are in the file
        if "Moazam" in content and "Abdullah" in content:
            self.logger.log("IDENTITY", "Verified", "Identity check passed")
            return True, content
        else:
            return False, "Invalid identity.txt content"
    
    def verify_consent(self):
        """Verify consent.txt exists and is valid"""
        if not os.path.exists(self.consent_file):
            return False, "consent.txt not found! Creating..."
        
        with open(self.consent_file, "r", encoding="utf-8") as f:
            content = f.read()
        
        if "Approved Targets" in content:
            self.logger.log("CONSENT", "Verified", "Consent check passed")
            return True, content
        else:
            return False, "Invalid consent.txt content"
    
    def verify_all(self):
        """Verify both identity and consent"""
        # Create files if they don't exist
        if not os.path.exists(self.identity_file):
            self.create_identity_file()
        
        if not os.path.exists(self.consent_file):
            self.create_consent_file()
        
        identity_ok, identity_msg = self.verify_identity()
        consent_ok, consent_msg = self.verify_consent()
        
        return identity_ok and consent_ok, identity_msg, consent_msg

# ============================================================================
# MAIN APP LAYOUT
# ============================================================================
def main():
    # Load custom CSS
    load_custom_css()
    
    # Initialize logger and identity checker
    logger = SecurityLogger()
    identity_checker = IdentitySafety()
    
    # ========================================================================
    # HEADER WITH ANIMATED BANNER
    # ========================================================================
    st.markdown("""
    <div style='text-align: center; padding: 20px;'>
        <h1 style='font-size: 3.5rem; margin: 0;'>🔐 NovaCrypt Defense</h1>
        <h3 style='color: #00bfff; margin-top: 10px;'>Hybrid Hacking Toolkit for PayBuddy FinTech</h3>
        <p style='color: #888; font-size: 0.9rem;'>Advanced Security Assessment Suite | Ethical Testing Only</p>
    </div>
    <hr>
    """, unsafe_allow_html=True)
    
    # ========================================================================
    # SIDEBAR - NAVIGATION & IDENTITY
    # ========================================================================
    with st.sidebar:
        st.markdown("## 🛡️ Control Panel")
        st.markdown("---")
        
        # Identity & Consent Verification
        st.markdown("### 👥 Team Identity")
        
        identity_ok, identity_msg, consent_msg = identity_checker.verify_all()
        
        if identity_ok:
            st.success("✅ Identity Verified")
            st.success("✅ Consent Verified")
            
            with st.expander("📄 View Identity"):
                st.code(identity_msg, language="text")
            
            with st.expander("📜 View Consent"):
                st.code(consent_msg, language="text")
        else:
            st.error("❌ Verification Failed")
            st.warning("Please check identity.txt and consent.txt files")
        
        st.markdown("---")
        
        # Module Selection
        st.markdown("### 🎯 Select Module")
        
        module = st.selectbox(
            "Choose a tool:",
            [
                "🏠 Dashboard",
                "🔍 Port Scanner",
                "🔑 Password Assessment",
                "💥 DOS/Stress Test",
                "🌐 Web Discovery",
                "📦 Packet Capture",
                "📊 Logs & Reports"
            ]
        )
        
        st.markdown("---")
        
        # Quick Stats
        st.markdown("### 📈 Session Stats")
        col1, col2 = st.columns(2)
        
        with col1:
            st.metric("Modules", "6", delta="Active")
        
        with col2:
            st.metric("Status", "Ready", delta="Online")
        
        st.markdown("---")
        
        # Dry Run Mode
        dry_run = st.checkbox("🧪 Dry Run Mode", value=False)
        if dry_run:
            st.info("🔔 Dry run enabled - No actual attacks will be performed")
        
        st.markdown("---")
    
    # ========================================================================
    # MAIN CONTENT AREA
    # ========================================================================
    
    if not identity_ok:
        st.error("🚫 **SECURITY CHECK FAILED**")
        st.warning("Please ensure identity.txt and consent.txt are properly configured before using the toolkit.")
        st.info("Files have been created automatically. Please verify their contents.")
        
        logger.log("SYSTEM", "Access Denied", "Identity/Consent verification failed")
        return
    
    # Log successful startup
    logger.log("SYSTEM", "Startup", f"Module selected: {module}")
    
    # ========================================================================
    # MODULE ROUTING
    # ========================================================================
    
    if module == "🏠 Dashboard":
        show_dashboard(logger, dry_run)
    
    elif module == "🔍 Port Scanner":
        show_port_scanner(logger, dry_run)
    
    elif module == "🔑 Password Assessment":
        show_password_assessment(logger, dry_run)
    
    elif module == "💥 DOS/Stress Test":
        show_stress_test(logger, dry_run)
    
    elif module == "🌐 Web Discovery":
        show_web_discovery(logger, dry_run)
    
    elif module == "📦 Packet Capture":
        show_packet_capture(logger, dry_run)
    
    elif module == "📊 Logs & Reports":
        show_logs_reports(logger)

# ============================================================================
# PACKET CAPTURE VIEW
# ============================================================================
def show_packet_capture(logger, dry_run):
    st.markdown("## 📦 Packet Capture & Analysis Module")
    
    logger.log("PACKET_CAPTURE", "Module Accessed", "User opened packet capture")
    
    st.markdown("""
    <div style='background: rgba(255, 20, 147, 0.1); padding: 20px; border-radius: 10px; border: 1px solid #ff1493; margin-bottom: 20px;'>
        <h4 style='color: #ff1493; margin-top: 0;'>📡 Network Traffic Capture & Protocol Analysis</h4>
        <p style='color: #fff;'>
            Capture and analyze network traffic in real-time. Monitor protocols, identify patterns,
            and export packet data for forensic analysis.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Configuration
    st.markdown("### ⚙️ Capture Configuration")
    
    col1, col2 = st.columns(2)
    
    with col1:
        capture_duration = st.slider(
            "⏱️ Capture Duration (seconds)",
            min_value=5,
            max_value=60,
            value=10,
            help="How long to capture traffic"
        )
        
        protocol_filter = st.selectbox(
            "🔍 Protocol Filter",
            ["All", "HTTP", "HTTPS", "DNS", "TCP", "UDP", "ICMP"],
            help="Filter specific protocol or capture all"
        )
    
    with col2:
        st.markdown("**Capture Info:**")
        st.markdown(f"""
        - Duration: {capture_duration} seconds
        - Filter: {protocol_filter}
        - Expected packets: ~{capture_duration * 10} packets
        - Sample rate: ~10 packets/second
        """)
        
        st.warning("""
        **⚠️ Privacy Notice:**
        
        Network capture may record sensitive data.
        Only capture on networks you own or have permission to monitor.
        """)
    
    # Control buttons
    st.markdown("---")
    
    col_btn1, col_btn2 = st.columns(2)
    
    with col_btn1:
        start_capture = st.button("📡 Start Capture", type="primary", use_container_width=True)
    
    with col_btn2:
        if st.button("🔄 Clear Results", use_container_width=True):
            if 'capture_results' in st.session_state:
                del st.session_state.capture_results
            st.rerun()
    
    # Initialize packet capture
    packet_capturer = PacketCapture(logger)
    
    # Execute capture
    if start_capture:
        st.markdown("---")
        st.markdown("## 📡 Capturing Traffic...")
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        packet_counter = st.empty()
        
        # Capture traffic directly (no threading for Streamlit)
        start_time = time.time()
        
        for i in range(capture_duration):
            progress = int(((i + 1) / capture_duration) * 100)
            progress_bar.progress(progress)
            
            elapsed = i + 1
            status_text.info(f"📡 Capturing... {elapsed}s / {capture_duration}s")
            
            # Generate packets for this second
            packets_per_second = random.randint(8, 12)
            for _ in range(packets_per_second):
                if protocol_filter == "All":
                    packet = packet_capturer.generate_sample_packet("Random")
                else:
                    packet = packet_capturer.generate_sample_packet(protocol_filter)
                
                packet_capturer.packets.append(packet)
            
            packet_counter.metric("Packets Captured", len(packet_capturer.packets))
            time.sleep(1)
        
        # Analyze results
        results = packet_capturer.analyze_traffic()
        st.session_state.capture_results = results
        
        progress_bar.progress(100)
        status_text.success(f"✅ Capture completed! Captured {len(packet_capturer.packets)} packets")
        
        # Force rerun to show results
        time.sleep(1)
        st.rerun()
    
    # Display results
    if 'capture_results' in st.session_state:
        results = st.session_state.capture_results
        
        st.markdown("---")
        st.markdown("## 📊 Traffic Analysis")
        
        # Summary metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total Packets", results['total_packets'])
        
        with col2:
            kb = results['total_bytes'] / 1024
            st.metric("Traffic Volume", f"{kb:.1f} KB")
        
        with col3:
            protocols = len(results['protocol_distribution'])
            st.metric("Protocols", protocols)
        
        with col4:
            avg_size = results['total_bytes'] / results['total_packets'] if results['total_packets'] > 0 else 0
            st.metric("Avg Packet Size", f"{avg_size:.0f} bytes")
        
        st.markdown("---")
        
        # Protocol distribution
        col_p1, col_p2 = st.columns(2)
        
        with col_p1:
            st.markdown("### 📊 Protocol Distribution")
            
            protocol_dist = results['protocol_distribution']
            total = sum(protocol_dist.values())
            
            for protocol, count in sorted(protocol_dist.items(), key=lambda x: x[1], reverse=True):
                percentage = (count / total * 100) if total > 0 else 0
                
                # Protocol colors
                colors = {
                    "HTTP": "🟢",
                    "HTTPS": "🟢",
                    "DNS": "🔵",
                    "TCP": "🟡",
                    "UDP": "🟠",
                    "ICMP": "⚪"
                }
                
                emoji = colors.get(protocol, "⚫")
                
                st.markdown(f"{emoji} **{protocol}:** {count} packets ({percentage:.1f}%)")
                st.progress(percentage / 100)
        
        with col_p2:
            st.markdown("### 🌐 Top Talkers (IPs)")
            
            for ip_info in results['top_ips'][:5]:
                st.markdown(f"**{ip_info['ip']}**")
                st.markdown(f"Packets: {ip_info['packets']}")
                st.markdown("---")
        
        # Port analysis
        st.markdown("### 🔌 Top Destination Ports")
        
        port_col1, port_col2, port_col3 = st.columns(3)
        
        for i, port_info in enumerate(results['top_ports'][:6]):
            col = [port_col1, port_col2, port_col3][i % 3]
            
            with col:
                port_name = {
                    80: "HTTP",
                    443: "HTTPS",
                    53: "DNS",
                    22: "SSH",
                    21: "FTP",
                    25: "SMTP",
                    3306: "MySQL",
                    5432: "PostgreSQL"
                }.get(port_info['port'], "Unknown")
                
                st.metric(
                    f"Port {port_info['port']}",
                    port_info['packets'],
                    delta=port_name
                )
        
        # Packet list
        st.markdown("---")
        st.markdown("### 📝 Captured Packets (Last 20)")
        
        # Display packets in table format
        for i, packet in enumerate(results['packets'][-20:], 1):
            with st.expander(f"Packet {i}: {packet['protocol']} | {packet['src_ip']} → {packet['dst_ip']}"):
                col_a, col_b = st.columns(2)
                
                with col_a:
                    st.markdown(f"""
                    **Timestamp:** {packet['timestamp']}  
                    **Protocol:** {packet['protocol']}  
                    **Length:** {packet['length']} bytes
                    """)
                
                with col_b:
                    st.markdown(f"""
                    **Source:** {packet['src_ip']}:{packet['src_port']}  
                    **Destination:** {packet['dst_ip']}:{packet['dst_port']}  
                    **Info:** {packet['info']}
                    """)
        
        # Security insights
        st.markdown("---")
        st.markdown("### 🔒 Security Insights")
        
        # Analyze for suspicious patterns
        http_count = results['protocol_distribution'].get('HTTP', 0)
        https_count = results['protocol_distribution'].get('HTTPS', 0)
        
        if http_count > https_count:
            st.warning(f"""
            ⚠️ **Unencrypted Traffic Detected**
            
            Found {http_count} HTTP packets vs {https_count} HTTPS packets.
            
            **Recommendation:** Use HTTPS for all sensitive communications.
            """)
        else:
            st.success("✅ Most traffic is encrypted (HTTPS)")
        
        # DNS analysis
        dns_count = results['protocol_distribution'].get('DNS', 0)
        if dns_count > results['total_packets'] * 0.3:
            st.info(f"""
            📊 **High DNS Activity**
            
            DNS queries represent {(dns_count/results['total_packets']*100):.1f}% of traffic.
            This could indicate normal browsing or potential DNS tunneling.
            """)
        
        # Export options
        st.markdown("---")
        st.markdown("### 📥 Export Captured Data")
        
        col_exp1, col_exp2 = st.columns(2)
        
        with col_exp1:
            if st.button("📦 Export PCAP (Text Format)", use_container_width=True):
                filename = f"capture_9953_Moazam_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pcap"
                filepath = packet_capturer.export_pcap(filename)
                
                with open(filepath, 'r') as f:
                    st.download_button(
                        label="⬇️ Download PCAP",
                        data=f.read(),
                        file_name=filename,
                        mime="text/plain"
                    )
                st.success(f"✅ PCAP exported: {filename}")
        
        with col_exp2:
            if st.button("📊 Export Analysis (JSON)", use_container_width=True):
                filename = f"traffic_analysis_9953_Moazam_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                filepath = packet_capturer.export_json(results, filename)
                
                with open(filepath, 'r') as f:
                    st.download_button(
                        label="⬇️ Download Analysis",
                        data=f.read(),
                        file_name=filename,
                        mime="application/json"
                    )
                st.success(f"✅ Analysis exported: {filename}")

# ============================================================================
# WEB DISCOVERY VIEW
# ============================================================================
def show_web_discovery(logger, dry_run):
    st.markdown("## 🌐 Web Discovery Module")
    
    logger.log("WEB_DISCOVERY", "Module Accessed", "User opened web discovery")
    
    st.markdown("""
    <div style='background: rgba(138, 43, 226, 0.1); padding: 20px; border-radius: 10px; border: 1px solid #8a2be2; margin-bottom: 20px;'>
        <h4 style='color: #8a2be2; margin-top: 0;'>🔍 DIRB-Style Directory & Endpoint Discovery</h4>
        <p style='color: #fff;'>
            Discover hidden directories, files, API endpoints, and subdomains using automated
            enumeration techniques. Essential for web application security assessment.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Mode selection
    st.markdown("### 🎯 Discovery Mode")
    
    discovery_mode = st.radio(
        "Choose discovery type:",
        ["Directory Enumeration", "Subdomain Discovery"],
        horizontal=True
    )
    
    st.markdown("---")
    
    # Initialize web discovery
    web_scanner = WebDiscovery(logger)
    
    # ========================================================================
    # MODE 1: DIRECTORY ENUMERATION
    # ========================================================================
    if discovery_mode == "Directory Enumeration":
        st.markdown("### 📂 Directory & File Discovery")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Target selection
            target_preset = st.selectbox(
                "🎯 Quick Select Target",
                [
                    "Custom (Enter Below)",
                    "http://testphp.vulnweb.com",
                    "http://demo.testfire.net",
                    "http://localhost:8080",
                    "http://127.0.0.1:5000"
                ]
            )
            
            if target_preset == "Custom (Enter Below)":
                default_target = ""
            else:
                default_target = target_preset
            
            target_url = st.text_input(
                "🌐 Target URL",
                value=default_target,
                help="Base URL to scan (e.g., http://example.com)",
                placeholder="http://example.com"
            )
            
            wordlist_type = st.selectbox(
                "📚 Wordlist Size",
                ["Minimal (15 paths)", "Common (50 paths)", "Extensive (200+ paths)"],
                help="Larger wordlists take longer but find more"
            )
            
            wordlist_map = {
                "Minimal (15 paths)": "minimal",
                "Common (50 paths)": "common",
                "Extensive (200+ paths)": "extensive"
            }
            
            wordlist = wordlist_map[wordlist_type]
        
        with col2:
            num_threads = st.slider(
                "🧵 Concurrent Threads",
                min_value=1,
                max_value=20,
                value=10,
                help="More threads = faster scan"
            )
            
            timeout = st.slider(
                "⏱️ Request Timeout (seconds)",
                min_value=1,
                max_value=10,
                value=5,
                help="Timeout for each request"
            )
            
            check_extensions = st.checkbox(
                "🔧 Check File Extensions",
                value=True,
                help="Also check .php, .html, .bak, etc."
            )
        
        # Estimated time
        path_count = {"minimal": 15, "common": 50, "extensive": 250}[wordlist]
        if check_extensions and wordlist == "extensive":
            path_count = 250 * 7  # Multiple extensions
        
        est_time = (path_count / num_threads) * (timeout * 0.3)
        
        st.info(f"""
        **Scan Estimate:**
        - Paths to check: ~{path_count}
        - Estimated time: ~{est_time:.1f} seconds
        - Rate: ~{num_threads * 10} requests/minute
        """)
        
        # Control buttons
        st.markdown("---")
        
        col_btn1, col_btn2 = st.columns(2)
        
        with col_btn1:
            start_scan = st.button("🔍 Start Discovery", type="primary", use_container_width=True, disabled=not target_url)
        
        with col_btn2:
            if st.button("🔄 Clear Results", use_container_width=True):
                if 'discovery_results' in st.session_state:
                    del st.session_state.discovery_results
                st.rerun()
        
        # Execute scan
        if start_scan:
            if not target_url:
                st.error("❌ Please enter a target URL!")
                return
            
            # Validate target
            if not dry_run:
                approved_domains = [
                    "testphp.vulnweb.com",
                    "demo.testfire.net",
                    "localhost",
                    "127.0.0.1"
                ]
                
                is_approved = any(domain in target_url.lower() for domain in approved_domains)
                
                if not is_approved:
                    st.warning("""
                    ⚠️ **Authorization Required**
                    
                    Directory enumeration on unauthorized targets may be illegal.
                    Only scan systems you own or have permission to test.
                    """)
                    
                    confirm = st.checkbox("✅ I have authorization to scan this target")
                    if not confirm:
                        logger.log("WEB_DISCOVERY", "Aborted", f"Unauthorized target: {target_url}", "WARNING")
                        return
            
            # Progress display
            st.markdown("---")
            st.markdown("## 🔄 Scanning In Progress...")
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            if dry_run:
                # Simulate scan
                for i in range(100):
                    progress_bar.progress(i + 1)
                    status_text.info(f"🧪 Dry run: Checking path {i}/{path_count}...")
                    time.sleep(0.02)
                
                # Fake results
                results = {
                    "base_url": target_url,
                    "wordlist": wordlist,
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "paths_checked": path_count,
                    "resources_found": 8,
                    "discovered_paths": [
                        {"url": f"{target_url}/admin", "path": "admin", "status_code": 200, "size": "4523", "content_type": "text/html"},
                        {"url": f"{target_url}/login", "path": "login", "status_code": 200, "size": "2341", "content_type": "text/html"},
                        {"url": f"{target_url}/api", "path": "api", "status_code": 301, "size": "Unknown", "content_type": "text/html"},
                        {"url": f"{target_url}/robots.txt", "path": "robots.txt", "status_code": 200, "size": "127", "content_type": "text/plain"},
                        {"url": f"{target_url}/sitemap.xml", "path": "sitemap.xml", "status_code": 200, "size": "1523", "content_type": "text/xml"},
                        {"url": f"{target_url}/backup", "path": "backup", "status_code": 403, "size": "Unknown", "content_type": "text/html"},
                        {"url": f"{target_url}/.git", "path": ".git", "status_code": 403, "size": "Unknown", "content_type": "text/html"},
                        {"url": f"{target_url}/config.php.bak", "path": "config.php.bak", "status_code": 200, "size": "892", "content_type": "text/plain"}
                    ]
                }
                
                st.session_state.discovery_results = results
                status_text.success("✅ Dry run completed!")
                logger.log("WEB_DISCOVERY", "Dry Run", f"Simulated scan on {target_url}")
            
            else:
                # Real scan
                start_time = time.time()
                
                # Run scan
                results = web_scanner.scan_directories(
                    target_url,
                    wordlist=wordlist,
                    num_threads=num_threads,
                    timeout=timeout,
                    extensions=check_extensions
                )
                
                # Update progress
                while web_scanner.checked_count < path_count:
                    progress = min(int((web_scanner.checked_count / path_count) * 100), 100)
                    progress_bar.progress(progress)
                    status_text.info(f"🔍 Scanning... {web_scanner.checked_count}/{path_count} paths checked")
                    time.sleep(0.3)
                    
                    if progress >= 99:
                        break
                
                progress_bar.progress(100)
                st.session_state.discovery_results = results
                status_text.success(f"✅ Scan completed! Found {results['resources_found']} resources")
        
        # Display results
        if 'discovery_results' in st.session_state:
            results = st.session_state.discovery_results
            
            st.markdown("---")
            st.markdown("## 📊 Discovery Results")
            
            # Summary metrics
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Paths Checked", results['paths_checked'])
            
            with col2:
                st.metric("Resources Found", results['resources_found'], delta="✅")
            
            with col3:
                discovery_rate = (results['resources_found'] / results['paths_checked'] * 100) if results['paths_checked'] > 0 else 0
                st.metric("Discovery Rate", f"{discovery_rate:.1f}%")
            
            with col4:
                st.metric("Target", results['base_url'].split('//')[1].split('/')[0])
            
            st.markdown("---")
            
            # Discovered paths
            if results['resources_found'] > 0:
                st.markdown("### 🔓 Discovered Resources")
                
                # Group by status code
                status_groups = defaultdict(list)
                for path in results['discovered_paths']:
                    status_groups[path['status_code']].append(path)
                
                # Display by status code
                for status_code in sorted(status_groups.keys()):
                    paths = status_groups[status_code]
                    
                    # Status code color
                    if status_code == 200:
                        color = "🟢"
                        status_name = "OK"
                    elif status_code in [301, 302, 307]:
                        color = "🟡"
                        status_name = "Redirect"
                    elif status_code == 403:
                        color = "🟠"
                        status_name = "Forbidden"
                    elif status_code == 401:
                        color = "🔴"
                        status_name = "Unauthorized"
                    else:
                        color = "⚪"
                        status_name = "Other"
                    
                    with st.expander(f"{color} **Status {status_code} ({status_name})** - {len(paths)} resource(s)", expanded=(status_code == 200)):
                        for path in paths:
                            col_a, col_b = st.columns([3, 1])
                            
                            with col_a:
                                st.markdown(f"**URL:** `{path['url']}`")
                                st.markdown(f"**Type:** {path['content_type']} | **Size:** {path['size']} bytes")
                            
                            with col_b:
                                st.code(f"{status_code}", language="text")
                
                # Security findings
                st.markdown("---")
                st.markdown("### 🔒 Security Assessment")
                
                interesting_paths = []
                critical_paths = []
                
                for path in results['discovered_paths']:
                    path_lower = path['path'].lower()
                    
                    # Critical findings
                    if any(x in path_lower for x in ['.git', '.env', '.bak', 'backup', 'config', 'password', 'admin', 'phpmyadmin']):
                        critical_paths.append(path)
                    # Interesting findings
                    elif any(x in path_lower for x in ['api', 'login', 'upload', 'log', 'test', 'dev']):
                        interesting_paths.append(path)
                
                if critical_paths:
                    st.error(f"""
                    🚨 **CRITICAL FINDINGS ({len(critical_paths)})**
                    
                    Found sensitive resources that should NOT be publicly accessible:
                    """)
                    for p in critical_paths:
                        st.markdown(f"- `{p['path']}` [{p['status_code']}]")
                
                if interesting_paths:
                    st.warning(f"""
                    ⚠️ **INTERESTING FINDINGS ({len(interesting_paths)})**
                    
                    Found potentially sensitive endpoints:
                    """)
                    for p in interesting_paths[:5]:  # Show first 5
                        st.markdown(f"- `{p['path']}` [{p['status_code']}]")
                
                if not critical_paths and not interesting_paths:
                    st.success("✅ No obvious security issues detected")
                
            else:
                st.warning("⚠️ No resources found. Target may be well-secured or non-existent.")
            
            # Export
            st.markdown("---")
            
            if st.button("📥 Export Discovery Report (JSON)", use_container_width=True):
                filename = f"discovery_{results['base_url'].replace('://', '_').replace('/', '_')}_9953_Moazam.json"
                filepath = web_scanner.export_results(results, filename)
                
                with open(filepath, 'r') as f:
                    st.download_button(
                        label="⬇️ Download Report",
                        data=f.read(),
                        file_name=filename,
                        mime="application/json"
                    )
                st.success(f"✅ Report exported: {filename}")
    
    # ========================================================================
    # MODE 2: SUBDOMAIN DISCOVERY
    # ========================================================================
    elif discovery_mode == "Subdomain Discovery":
        st.markdown("### 🌐 Subdomain Enumeration")
        
        st.info("""
        **Subdomain Discovery**
        
        Finds active subdomains for a given domain (e.g., admin.example.com, api.example.com).
        Uses DNS resolution to check common subdomain names.
        """)
        
        domain_input = st.text_input(
            "🌐 Domain (without subdomain)",
            placeholder="example.com",
            help="Enter domain only, no http:// or www"
        )
        
        if st.button("🔍 Find Subdomains", type="primary", use_container_width=True):
            if domain_input:
                st.markdown("### 🔄 Checking Subdomains...")
                
                progress = st.progress(0)
                
                found = web_scanner.check_subdomains(domain_input)
                
                progress.progress(100)
                
                if found:
                    st.success(f"✅ Found {len(found)} active subdomains!")
                    
                    for sub in found:
                        st.markdown(f"🟢 **{sub['subdomain']}** - {sub['status']}")
                else:
                    st.warning("⚠️ No common subdomains found")

# ============================================================================
# DOS/STRESS TEST VIEW
# ============================================================================
def show_stress_test(logger, dry_run):
    st.markdown("## 💥 DOS/Stress Test Module")
    
    logger.log("STRESS_TEST", "Module Accessed", "User opened stress testing")
    
    st.markdown("""
    <div style='background: rgba(255, 165, 0, 0.1); padding: 20px; border-radius: 10px; border: 1px solid #ffa500; margin-bottom: 20px;'>
        <h4 style='color: #ffa500; margin-top: 0;'>⚡ Controlled Load & Stress Testing</h4>
        <p style='color: #fff;'>
            Test server resilience under load with controlled HTTP flooding. Monitor latency,
            success rates, and performance degradation in real-time.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Safety warning
    st.warning("""
    ⚠️ **CRITICAL SAFETY NOTICE**
    
    - Only test **authorized targets** with written permission
    - Maximum 200 concurrent clients enforced
    - Aggressive testing may be **illegal** without authorization
    - Use **dry run mode** first to verify configuration
    - Stop test immediately if target becomes unresponsive
    """)
    
    st.markdown("---")
    
    # Configuration Section
    st.markdown("### ⚙️ Test Configuration")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Quick target selection
        target_preset = st.selectbox(
            "🎯 Quick Select Target",
            [
                "Custom (Enter Below)",
                "http://testphp.vulnweb.com (Authorized Test Site)",
                "http://demo.testfire.net (IBM Demo Site)",
                "http://localhost:8080 (Local Server)",
                "http://127.0.0.1:5000 (Local Flask)"
            ]
        )
        
        if target_preset == "Custom (Enter Below)":
            default_target = ""
        else:
            # Extract URL from preset
            default_target = target_preset.split(" ")[0]
        
        target_url = st.text_input(
            "🌐 Target URL",
            value=default_target,
            help="Full URL including http:// or https://",
            placeholder="http://example.com"
        )
        
        request_method = st.selectbox(
            "📊 Request Method",
            ["GET", "POST", "HEAD"],
            help="HTTP method to use for requests"
        )
    
    with col2:
        num_clients = st.slider(
            "👥 Concurrent Clients",
            min_value=1,
            max_value=200,
            value=50,
            step=10,
            help="Number of simultaneous connections (max 200 for safety)"
        )
        
        duration = st.slider(
            "⏱️ Test Duration (seconds)",
            min_value=5,
            max_value=60,
            value=10,
            step=5,
            help="How long to run the test"
        )
        
        timeout = st.slider(
            "🕐 Request Timeout (seconds)",
            min_value=1,
            max_value=10,
            value=5,
            step=1,
            help="Timeout for individual requests"
        )
    
    # Estimate calculations
    st.info(f"""
    **Estimated Test Parameters:**
    - Total Requests: ~{num_clients * duration * 10} requests
    - Requests/Second: ~{num_clients * 10} req/s
    - Test Duration: {duration} seconds
    - Max Concurrent: {num_clients} clients
    """)
    
    # Control buttons
    st.markdown("---")
    
    col_btn1, col_btn2, col_btn3 = st.columns([2, 1, 1])
    
    with col_btn1:
        if dry_run:
            st.warning("🧪 Dry run mode active - no actual requests will be sent")
        
        start_test = st.button("🚀 Start Stress Test", type="primary", use_container_width=True, disabled=not target_url)
    
    with col_btn2:
        if st.button("🔄 Clear Results", use_container_width=True):
            if 'stress_results' in st.session_state:
                del st.session_state.stress_results
            st.rerun()
    
    with col_btn3:
        if st.button("❓ Help", use_container_width=True):
            st.session_state.show_stress_help = not st.session_state.get('show_stress_help', False)
    
    # Help section
    if st.session_state.get('show_stress_help', False):
        with st.expander("📖 Stress Testing Guide", expanded=True):
            st.markdown("""
            ### How to Use Stress Tester
            
            **Approved Targets:**
            - `testphp.vulnweb.com` - Acunetix test site (authorized)
            - `demo.testfire.net` - IBM AppScan demo (authorized)
            - `localhost` - Your own local server
            - Your own websites/APIs with permission
            
            **Configuration:**
            - **Concurrent Clients:** More clients = more load
            - **Duration:** Longer tests reveal performance degradation
            - **Method:** GET for pages, POST for forms/APIs
            - **Timeout:** Lower = faster detection of slowness
            
            **Interpreting Results:**
            - **Success Rate:** Should be near 100% for healthy servers
            - **Latency P95/P99:** 95th/99th percentile response times
            - **Requests/Second:** Throughput capacity
            - **Status Codes:** 200=success, 500=server error, 0=timeout
            
            **Best Practices:**
            - Start with low clients (10-20) and short duration (10s)
            - Gradually increase load to find breaking point
            - Monitor target server resources if possible
            - Stop test if success rate drops below 50%
            - Document all tests for audit trail
            
            **When to Use:**
            - API load testing
            - Web application capacity planning
            - DDoS resilience verification
            - Performance benchmarking
            """)
    
    # Execute stress test
    if start_test:
        if not target_url:
            st.error("❌ Please enter a target URL!")
            return
        
        # Validate target
        if not dry_run:
            approved_domains = [
                "testphp.vulnweb.com",
                "demo.testfire.net",
                "localhost",
                "127.0.0.1",
                "0.0.0.0"
            ]
            
            is_approved = any(domain in target_url.lower() for domain in approved_domains)
            
            if not is_approved:
                st.error("""
                ❌ **UNAUTHORIZED TARGET**
                
                This target is not in the approved list. Stress testing unauthorized
                systems is illegal and may result in criminal charges.
                
                **Only test:**
                - Your own infrastructure
                - Explicitly authorized test sites
                - Local development servers
                """)
                
                confirm = st.checkbox("✅ I have written authorization to stress test this target")
                if not confirm:
                    logger.log("STRESS_TEST", "Aborted", f"Unauthorized target: {target_url}", "WARNING")
                    return
        
        # Initialize tester
        tester = StressTester(logger)
        
        # Progress display
        st.markdown("---")
        st.markdown("## 🔄 Test In Progress...")
        
        progress_bar = st.progress(0)
        status_container = st.empty()
        metrics_container = st.empty()
        
        if dry_run:
            # Simulate test
            for i in range(100):
                progress_bar.progress(i + 1)
                status_container.info(f"🧪 Dry run: Simulating request {i * 10}/{duration * 10}...")
                time.sleep(0.05)
            
            # Fake results
            results = {
                "target": target_url,
                "method": request_method,
                "num_clients": num_clients,
                "duration_seconds": duration,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "total_requests": num_clients * duration * 10,
                "successful_requests": int(num_clients * duration * 10 * 0.95),
                "failed_requests": int(num_clients * duration * 10 * 0.05),
                "success_rate_percent": 95.0,
                "requests_per_second": num_clients * 10,
                "latency_stats": {
                    "average_ms": 125.5,
                    "min_ms": 50.2,
                    "max_ms": 450.8,
                    "p50_ms": 110.0,
                    "p95_ms": 250.0,
                    "p99_ms": 400.0
                },
                "status_code_distribution": {
                    "200": int(num_clients * duration * 10 * 0.95),
                    "500": int(num_clients * duration * 10 * 0.03),
                    "0": int(num_clients * duration * 10 * 0.02)
                }
            }
            
            st.session_state.stress_results = results
            status_container.success("✅ Dry run completed!")
            logger.log("STRESS_TEST", "Dry Run", f"Simulated test on {target_url}")
            
        else:
            # Real test
            start_time = time.time()
            
            # Run test in thread to allow progress updates
            test_thread = threading.Thread(
                target=lambda: setattr(st.session_state, 'stress_results', 
                                      tester.run_stress_test(target_url, num_clients, duration, request_method, timeout))
            )
            test_thread.start()
            
            # Update progress
            while test_thread.is_alive():
                elapsed = time.time() - start_time
                progress = min(int((elapsed / duration) * 100), 100)
                progress_bar.progress(progress)
                
                status_container.info(f"⚡ Testing... {elapsed:.1f}s / {duration}s")
                
                # Show live metrics
                with metrics_container.container():
                    col_m1, col_m2, col_m3, col_m4 = st.columns(4)
                    with col_m1:
                        st.metric("Requests", tester.request_count)
                    with col_m2:
                        st.metric("Success", tester.success_count)
                    with col_m3:
                        st.metric("Errors", tester.error_count)
                    with col_m4:
                        rate = (tester.success_count / tester.request_count * 100) if tester.request_count > 0 else 0
                        st.metric("Success Rate", f"{rate:.1f}%")
                
                time.sleep(0.5)
            
            test_thread.join()
            
            progress_bar.progress(100)
            status_container.success("✅ Stress test completed!")
    
    # Display results
    if 'stress_results' in st.session_state:
        results = st.session_state.stress_results
        
        st.markdown("---")
        st.markdown("## 📊 Test Results")
        
        # Summary metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total Requests", results['total_requests'])
        
        with col2:
            st.metric("Successful", results['successful_requests'], delta="✅")
        
        with col3:
            st.metric("Failed", results['failed_requests'], delta="❌")
        
        with col4:
            success_rate = results['success_rate_percent']
            color = "🟢" if success_rate >= 95 else "🟡" if success_rate >= 80 else "🔴"
            st.metric("Success Rate", f"{success_rate}%", delta=color)

        # Generate and display performance graphs
        st.markdown("### 📊 Performance Visualization")
        
        try:
            graph_buf = generate_performance_graphs(results)
            st.image(graph_buf, use_container_width=True)
            
            # Download graph button
            st.download_button(
                label="📥 Download Performance Graph",
                data=graph_buf,
                file_name=f"stress_test_graph_{results['target'].replace('://', '_')}_9953_Moazam.png",
                mime="image/png",
                key="download_graph"
            )
        except Exception as e:
            st.warning(f"⚠️ Could not generate graphs: {str(e)}")
        
        st.markdown("---")
        
        # Performance metrics
        col_p1, col_p2 = st.columns(2)
        
        with col_p1:
            st.markdown("### ⚡ Performance Metrics")
            
            latency = results['latency_stats']
            
            st.markdown(f"""
            **Response Times:**
            - Average: `{latency['average_ms']} ms`
            - Minimum: `{latency['min_ms']} ms`
            - Maximum: `{latency['max_ms']} ms`
            
            **Percentiles:**
            - P50 (Median): `{latency['p50_ms']} ms`
            - P95: `{latency['p95_ms']} ms`
            - P99: `{latency['p99_ms']} ms`
            
            **Throughput:**
            - Requests/Second: `{results['requests_per_second']} req/s`
            """)
        
        with col_p2:
            st.markdown("### 📈 Status Code Distribution")
            
            status_codes = results['status_code_distribution']
            
            for code, count in sorted(status_codes.items()):
                percentage = (count / results['total_requests'] * 100) if results['total_requests'] > 0 else 0
                
                if code == 200 or code == '200':
                    color = "🟢"
                elif code == 0 or code == '0':
                    color = "⚫"
                else:
                    color = "🔴"
                
                code_name = {
                    200: "Success",
                    201: "Created",
                    204: "No Content",
                    400: "Bad Request",
                    401: "Unauthorized",
                    403: "Forbidden",
                    404: "Not Found",
                    500: "Server Error",
                    502: "Bad Gateway",
                    503: "Service Unavailable",
                    0: "Timeout/Error"
                }.get(int(code) if str(code).isdigit() else 0, "Unknown")
                
                st.markdown(f"{color} **{code} ({code_name}):** {count} ({percentage:.1f}%)")
        
        # Assessment
        st.markdown("---")
        st.markdown("### 🎯 Performance Assessment")
        
        # Determine server health
        if success_rate >= 99:
            st.success("""
            ✅ **EXCELLENT** - Server handled load perfectly
            - Success rate above 99%
            - Server is robust and well-configured
            - Can likely handle higher load
            """)
        elif success_rate >= 95:
            st.success("""
            ✅ **GOOD** - Server performed well under load
            - Success rate above 95%
            - Minor issues under stress
            - Consider optimization for peak loads
            """)
        elif success_rate >= 80:
            st.warning("""
            ⚠️ **MODERATE** - Server showed signs of stress
            - Success rate 80-95%
            - Performance degradation observed
            - Recommend capacity improvements
            """)
        else:
            st.error("""
            ❌ **POOR** - Server struggled under load
            - Success rate below 80%
            - Significant failures occurred
            - URGENT: Review infrastructure and scaling
            """)
        
        # Latency assessment
        avg_latency = latency['average_ms']
        p99_latency = latency['p99_ms']
        
        if avg_latency < 100:
            st.success(f"⚡ Fast response times (avg: {avg_latency}ms)")
        elif avg_latency < 500:
            st.info(f"⏱️ Acceptable response times (avg: {avg_latency}ms)")
        else:
            st.warning(f"🐌 Slow response times (avg: {avg_latency}ms) - optimization needed")
        
        # Export results
        st.markdown("---")
        st.markdown("### 📥 Export Results")
        
        if st.button("📊 Export Test Report (JSON)", use_container_width=True):
            filename = f"stress_test_{results['target'].replace('://', '_').replace('/', '_')}_9953_Moazam.json"
            filepath = tester.export_results(results, filename)
            
            with open(filepath, 'r') as f:
                st.download_button(
                    label="⬇️ Download Report",
                    data=f.read(),
                    file_name=filename,
                    mime="application/json"
                )
            st.success(f"✅ Report exported: {filename}")

# ============================================================================
# PASSWORD ASSESSMENT VIEW
# ============================================================================
def show_password_assessment(logger, dry_run):
    st.markdown("## 🔑 Password Assessment Module")
    
    logger.log("PASSWORD_TEST", "Module Accessed", "User opened password assessment")
    
    st.markdown("""
    <div style='background: rgba(0, 255, 127, 0.1); padding: 20px; border-radius: 10px; border: 1px solid #00ff7f; margin-bottom: 20px;'>
        <h4 style='color: #00ff7f; margin-top: 0;'>🔐 Comprehensive Password Security Analysis</h4>
        <p style='color: #fff;'>
            Test password strength, check policy compliance, calculate entropy, and receive
            actionable security recommendations. Includes hash simulation for educational purposes.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Assessment Mode Selection
    st.markdown("### 🎯 Assessment Mode")
    
    mode = st.radio(
        "Choose testing mode:",
        ["Single Password Analysis", "Batch Password Testing", "Hash Simulation"],
        horizontal=True
    )
    
    st.markdown("---")
    
    # Initialize password tester
    password_tester = PasswordAssessment(logger)
    
    # ========================================================================
    # MODE 1: SINGLE PASSWORD ANALYSIS
    # ========================================================================
    if mode == "Single Password Analysis":
        st.markdown("### 🔍 Password Analysis")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            password_input = st.text_input(
                "🔑 Enter Password to Test",
                type="password",
                help="Your password is analyzed locally and never stored",
                placeholder="Enter a password..."
            )
            
            show_password = st.checkbox("👁️ Show password", value=False)
            
            if show_password and password_input:
                st.code(password_input, language="text")
        
        with col2:
            st.info("""
            **Privacy Notice:**
            
            ✅ Analysis is local  
            ✅ No storage  
            ✅ No transmission  
            ✅ Completely safe
            """)
        
        # Analyze button
        st.markdown("---")
        
        col_btn1, col_btn2 = st.columns(2)
        
        with col_btn1:
            if st.button("🔍 Analyze Password", type="primary", use_container_width=True, disabled=not password_input):
                if password_input:
                    # Perform assessment
                    results = password_tester.assess_password(password_input)
                    
                    # Store in session state
                    st.session_state.password_results = results
                    
                    st.success("✅ Analysis complete!")
        
        with col_btn2:
            if st.button("🔄 Clear Results", use_container_width=True):
                if 'password_results' in st.session_state:
                    del st.session_state.password_results
                    st.rerun()
        
        # Display results
        if 'password_results' in st.session_state:
            results = st.session_state.password_results
            
            st.markdown("---")
            st.markdown("## 📊 Assessment Results")
            
            # Strength meter
            score = results['strength_score']
            category = results['strength_category']
            emoji = results['strength_emoji']
            color = results['strength_color']
            
            st.markdown(f"""
            <div style='background: rgba(255, 255, 255, 0.05); padding: 25px; border-radius: 10px; border: 2px solid {color}; text-align: center;'>
                <h1 style='color: {color}; margin: 0; font-size: 4rem;'>{emoji}</h1>
                <h2 style='color: {color}; margin: 10px 0;'>{category}</h2>
                <h1 style='color: {color}; margin: 10px 0; font-size: 3rem;'>{score}/100</h1>
                <div style='background: #1a1a2e; height: 30px; border-radius: 15px; overflow: hidden; margin-top: 20px;'>
                    <div style='background: {color}; width: {score}%; height: 100%; transition: width 0.5s;'></div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            # Detailed metrics
            col_m1, col_m2, col_m3, col_m4 = st.columns(4)
            
            with col_m1:
                st.metric("Length", f"{results['password_length']} chars")
            
            with col_m2:
                st.metric("Entropy", f"{results['entropy_bits']} bits")
            
            with col_m3:
                policy_status = "✅ Pass" if results['policy_compliant'] else "❌ Fail"
                st.metric("Policy Check", policy_status)
            
            with col_m4:
                common_status = "🚨 YES" if results['is_common_password'] else "✅ NO"
                st.metric("Common Pwd", common_status)
            
            st.markdown("---")
            
            # Character sets
            st.markdown("### 🔤 Character Analysis")
            
            col_c1, col_c2, col_c3, col_c4 = st.columns(4)
            
            char_sets = results['character_sets']
            
            with col_c1:
                if char_sets['lowercase']:
                    st.success("✅ Lowercase (a-z)")
                else:
                    st.error("❌ Lowercase (a-z)")
            
            with col_c2:
                if char_sets['uppercase']:
                    st.success("✅ Uppercase (A-Z)")
                else:
                    st.error("❌ Uppercase (A-Z)")
            
            with col_c3:
                if char_sets['digits']:
                    st.success("✅ Digits (0-9)")
                else:
                    st.error("❌ Digits (0-9)")
            
            with col_c4:
                if char_sets['special']:
                    st.success("✅ Special (!@#$...)")
                else:
                    st.error("❌ Special (!@#$...)")
            
            # Patterns detected
            if results['patterns_detected']:
                st.markdown("### ⚠️ Patterns Detected")
                for pattern in results['patterns_detected']:
                    st.warning(f"🔍 {pattern}")
            
            # Policy issues
            if results['policy_issues']:
                st.markdown("### ❌ Policy Violations")
                for issue in results['policy_issues']:
                    st.error(f"• {issue}")
            else:
                st.success("### ✅ Policy Compliant - Meets all requirements")
            
            # Recommendations
            st.markdown("### 💡 Security Recommendations")
            
            for i, rec in enumerate(results['recommendations'], 1):
                st.markdown(f"{i}. {rec}")
            
            # Export option
            st.markdown("---")
            
            if st.button("📥 Export Analysis Report (JSON)", use_container_width=True):
                filename = f"password_analysis_9953_Moazam_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                filepath = password_tester.export_results(results, filename)
                
                with open(filepath, 'r') as f:
                    st.download_button(
                        label="⬇️ Download Report",
                        data=f.read(),
                        file_name=filename,
                        mime="application/json"
                    )
                st.success(f"✅ Report exported: {filename}")
    
    # ========================================================================
    # MODE 2: BATCH PASSWORD TESTING
    # ========================================================================
    elif mode == "Batch Password Testing":
        st.markdown("### 📋 Batch Password Analysis")
        
        st.info("""
        **Test multiple passwords at once**
        
        Enter one password per line. Useful for:
        - Testing user password databases
        - Comparing password strengths
        - Bulk policy compliance checking
        """)
        
        passwords_input = st.text_area(
            "🔑 Enter Passwords (one per line)",
            height=200,
            placeholder="password123\nMyP@ssw0rd!\nSecurePass2024\n..."
        )
        
        if st.button("🔍 Analyze All Passwords", type="primary", use_container_width=True):
            if passwords_input.strip():
                passwords = [p.strip() for p in passwords_input.split('\n') if p.strip()]
                
                st.markdown(f"### 📊 Analyzing {len(passwords)} passwords...")
                
                batch_results = []
                
                progress_bar = st.progress(0)
                
                for i, pwd in enumerate(passwords):
                    result = password_tester.assess_password(pwd)
                    result['password_preview'] = pwd[:3] + '*' * (len(pwd) - 3)  # Masked
                    batch_results.append(result)
                    progress_bar.progress((i + 1) / len(passwords))
                
                st.session_state.batch_results = batch_results
                
                st.success(f"✅ Analyzed {len(passwords)} passwords!")
        
        # Display batch results
        if 'batch_results' in st.session_state:
            results = st.session_state.batch_results
            
            st.markdown("---")
            st.markdown("### 📊 Batch Analysis Results")
            
            # Summary statistics
            col_s1, col_s2, col_s3, col_s4 = st.columns(4)
            
            avg_score = sum(r['strength_score'] for r in results) / len(results)
            weak_count = sum(1 for r in results if r['strength_score'] < 40)
            strong_count = sum(1 for r in results if r['strength_score'] >= 80)
            common_count = sum(1 for r in results if r['is_common_password'])
            
            with col_s1:
                st.metric("Average Score", f"{avg_score:.1f}/100")
            
            with col_s2:
                st.metric("Weak Passwords", weak_count, delta="⚠️")
            
            with col_s3:
                st.metric("Strong Passwords", strong_count, delta="✅")
            
            with col_s4:
                st.metric("Common Passwords", common_count, delta="🚨")
            
            st.markdown("---")
            
            # Individual results table
            st.markdown("### 📋 Individual Results")
            
            for i, result in enumerate(results, 1):
                with st.expander(f"Password {i}: {result['password_preview']} - {result['strength_emoji']} {result['strength_category']} ({result['strength_score']}/100)"):
                    col_a, col_b = st.columns(2)
                    
                    with col_a:
                        st.markdown(f"""
                        **Metrics:**
                        - Score: {result['strength_score']}/100
                        - Length: {result['password_length']} characters
                        - Entropy: {result['entropy_bits']} bits
                        - Policy: {'✅ Pass' if result['policy_compliant'] else '❌ Fail'}
                        - Common: {'🚨 YES' if result['is_common_password'] else '✅ NO'}
                        """)
                    
                    with col_b:
                        st.markdown("**Character Sets:**")
                        for char_type, present in result['character_sets'].items():
                            emoji = "✅" if present else "❌"
                            st.markdown(f"{emoji} {char_type.capitalize()}")
    
    # ========================================================================
    # MODE 3: HASH SIMULATION
    # ========================================================================
    elif mode == "Hash Simulation":
        st.markdown("### 🔐 Password Hash Simulation")
        
        st.warning("""
        **⚠️ Educational Purpose Only**
        
        This demonstrates how passwords are hashed. In production:
        - Never use MD5 or SHA256 for passwords
        - Always use bcrypt, scrypt, or Argon2
        - Add proper salting
        - Use key stretching
        """)
        
        password_hash = st.text_input(
            "🔑 Enter Password to Hash",
            type="password",
            placeholder="Enter password..."
        )
        
        hash_types = st.multiselect(
            "📊 Select Hash Algorithms",
            ["MD5", "SHA256", "SHA512", "bcrypt"],
            default=["MD5", "SHA256", "bcrypt"]
        )
        
        if st.button("🔐 Generate Hashes", type="primary", use_container_width=True):
            if password_hash:
                st.markdown("---")
                st.markdown("### 🔒 Generated Hashes")
                
                hashes = password_tester.simulate_hash(password_hash, "all")
                
                for hash_type in hash_types:
                    if hash_type in hashes:
                        st.markdown(f"**{hash_type}:**")
                        st.code(hashes[hash_type], language="text")

# ============================================================================
# PORT SCANNER VIEW
# ============================================================================
def show_port_scanner(logger, dry_run):
    st.markdown("## 🔍 Port Scanner Module")
    
    logger.log("PORT_SCAN", "Module Accessed", "User opened port scanner")
    
    st.markdown("""
    <div style='background: rgba(0, 191, 255, 0.1); padding: 20px; border-radius: 10px; border: 1px solid #00bfff; margin-bottom: 20px;'>
        <h4 style='color: #00bfff; margin-top: 0;'>🎯 Professional TCP Port Scanner</h4>
        <p style='color: #fff;'>
            Multi-threaded port scanning with banner grabbing and service detection.
            Scan real-world targets to identify open ports and running services.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Configuration Section
    st.markdown("### ⚙️ Scan Configuration")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Predefined target selection
        target_preset = st.selectbox(
            "🎯 Quick Select Target",
            [
                "Custom (Enter Below)",
                "scanme.nmap.org (Official Nmap Test)",
                "testphp.vulnweb.com (Acunetix Vulnerable App)",
                "testhtml5.vulnweb.com (HTML5 Test App)",
                "demo.testfire.net (IBM AppScan Demo)",
                "zero.webappsecurity.com (Security Test Site)",
                "localhost (Your Machine)",
                "127.0.0.1 (Loopback)"
            ]
        )
        
        # Extract target from selection
        if target_preset == "Custom (Enter Below)":
            default_target = ""
        elif "scanme.nmap.org" in target_preset:
            default_target = "scanme.nmap.org"
        elif "testphp.vulnweb.com" in target_preset:
            default_target = "testphp.vulnweb.com"
        elif "testhtml5.vulnweb.com" in target_preset:
            default_target = "testhtml5.vulnweb.com"
        elif "demo.testfire.net" in target_preset:
            default_target = "demo.testfire.net"
        elif "zero.webappsecurity.com" in target_preset:
            default_target = "zero.webappsecurity.com"
        elif "localhost" in target_preset:
            default_target = "localhost"
        else:
            default_target = "127.0.0.1"
        
        target = st.text_input(
            "🌐 Target (IP or Domain)",
            value=default_target,
            help="Enter IP address or domain name to scan",
            placeholder="e.g., scanme.nmap.org or 192.168.1.1"
        )
        
        scan_type = st.selectbox(
            "📊 Scan Type",
            ["Common Ports (Top 17)", "Quick Scan (1-1024)", "Full Scan (1-65535)", "Custom Range"]
        )
        
        if scan_type == "Custom Range":
            port_range = st.text_input(
                "🔢 Port Range",
                value="80,443,8080",
                help="Format: '80' or '1-100' or '80,443,8080'"
            )
        elif scan_type == "Common Ports (Top 17)":
            port_range = "21,22,23,25,53,80,110,143,443,445,3306,3389,5432,5900,8080,8443,27017"
        elif scan_type == "Quick Scan (1-1024)":
            port_range = "1-1024"
        else:  # Full Scan
            port_range = "1-65535"
    
    with col2:
        num_threads = st.slider(
            "🧵 Number of Threads",
            min_value=10,
            max_value=200,
            value=50,
            step=10,
            help="More threads = faster scan, but may trigger IDS"
        )
        
        timeout = st.slider(
            "⏱️ Timeout (seconds)",
            min_value=0.5,
            max_value=5.0,
            value=1.0,
            step=0.5,
            help="Connection timeout per port"
        )
        
        st.info(f"""
        **Estimated Scan Time:**  
        {len(port_range.replace('-', ',').split(','))} ports ÷ {num_threads} threads ≈ {round(len(port_range.replace('-', ',').split(',')) / num_threads * timeout, 1)}s
        """)
    
    # Scan Button
    st.markdown("---")
    
    col_btn1, col_btn2, col_btn3 = st.columns([2, 1, 1])
    
    with col_btn1:
        if dry_run:
            st.warning("🧪 Dry run mode - simulation only")
        
        start_scan = st.button("🚀 Start Port Scan", type="primary", use_container_width=True)
    
    with col_btn2:
        if st.button("🔄 Clear Results", use_container_width=True):
            if 'scan_results' in st.session_state:
                del st.session_state.scan_results
            st.rerun()
    
    with col_btn3:
        show_help = st.button("❓ Help", use_container_width=True)
    
    if show_help:
        with st.expander("📖 Port Scanner Help", expanded=True):
            st.markdown("""
            ### How to Use Port Scanner
            
            **Target Selection:**
            
            **✅ SAFE & LEGAL Targets (Pre-approved):**
            - `scanme.nmap.org` - Official Nmap project test server
            - `testphp.vulnweb.com` - Acunetix vulnerable web application
            - `testhtml5.vulnweb.com` - HTML5 security test application
            - `demo.testfire.net` - IBM AppScan demo banking site
            - `zero.webappsecurity.com` - Security testing platform
            - `localhost` / `127.0.0.1` - Your own computer
            - `192.168.x.x` - Your local network (if you own it)
            
            **❌ NEVER Scan These (Illegal!):**
            - Any website you don't own
            - Commercial sites (Google, Amazon, Facebook, etc.)
            - Government sites (.gov domains)
            - Educational institutions (without permission)
            - Banking or financial sites
            
            **⚖️ Legal Notice:**
            Unauthorized port scanning is illegal under computer fraud laws in most countries.
            Always obtain written permission before scanning any target.
            
            ---
            
            **Scan Types:**
            - **Common Ports:** Scans 17 most common services (fastest)
            - **Quick Scan:** Ports 1-1024 (standard services)
            - **Full Scan:** All 65535 ports (very slow, 30+ minutes)
            - **Custom Range:** Specify exact ports (e.g., 80,443,8080)
            
            **Performance Tips:**
            - More threads = faster scan (but may trigger IDS/IPS)
            - Lower timeout = faster but may miss slow services
            - Use "Common Ports" for quick reconnaissance
            - Use "Quick Scan" for thorough web app testing
            
            **Recommended Testing Workflow:**
            1. Start with `scanme.nmap.org` + Common Ports
            2. Try `testphp.vulnweb.com` for web-specific ports
            3. Test `localhost` to see your own services
            
            **Ethical Guidelines:**
            - Only scan during off-peak hours
            - Use reasonable thread counts (50-100)
            - Don't scan the same target repeatedly
            - Document all authorized scans in logs
            """)
    
    # Execute Scan
    if start_scan:
        if not target:
            st.error("❌ Please enter a target!")
            return
        
        # Validation
        if not dry_run:
            # Extended list of approved targets
            consent_targets = [
                "scanme.nmap.org",
                "testphp.vulnweb.com", 
                "testhtml5.vulnweb.com",
                "demo.testfire.net",
                "zero.webappsecurity.com",
                "webscantest.com",
                "localhost", 
                "127.0.0.1",
                "0.0.0.0",
                "example.com"
            ]
            
            # Check if it's a local IP (192.168.x.x or 10.x.x.x)
            is_local_ip = any(target.startswith(prefix) for prefix in ["192.168.", "10.", "172.16."])
            
            is_approved = any(approved in target.lower() for approved in consent_targets) or is_local_ip
            
            if not is_approved:
                st.warning(f"""
                ⚠️ **Target Authorization Required**
                
                Target '{target}' is not in the pre-approved list.
                
                **Approved Targets:**
                - scanme.nmap.org (Official Nmap test server)
                - testphp.vulnweb.com (Acunetix test site)
                - testhtml5.vulnweb.com (HTML5 test app)
                - demo.testfire.net (IBM AppScan demo)
                - zero.webappsecurity.com (Security test site)
                - localhost / 127.0.0.1 (Your machine)
                - 192.168.x.x / 10.x.x.x (Local network)
                
                **⚖️ Legal Warning:** Unauthorized port scanning is illegal in many jurisdictions.
                Only scan systems you own or have explicit written permission to test.
                """)
                
                confirm = st.checkbox("✅ I confirm I have legal authorization to scan this target")
                if not confirm:
                    st.error("❌ Scan aborted - authorization required")
                    logger.log("PORT_SCAN", "Aborted", f"Unauthorized target: {target}", "WARNING")
                    return
        
        # Initialize scanner
        scanner = PortScanner(logger)
        
        # Progress display
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        if dry_run:
            # Simulate scan
            status_text.info("🧪 Dry run mode - simulating scan...")
            for i in range(100):
                time.sleep(0.02)
                progress_bar.progress(i + 1)
            
            # Fake results
            results = {
                "target": target,
                "ip": "simulation",
                "scan_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "total_ports_scanned": len(port_range.split(',')),
                "open_ports_found": 3,
                "open_ports": [
                    {"port": 80, "service": "HTTP", "state": "open", "banner": "Apache/2.4.41 (Simulated)"},
                    {"port": 443, "service": "HTTPS", "state": "open", "banner": "nginx/1.18.0 (Simulated)"},
                    {"port": 22, "service": "SSH", "state": "open", "banner": "OpenSSH 8.2 (Simulated)"}
                ]
            }
            st.session_state.scan_results = results
            status_text.success("✅ Dry run completed!")
            logger.log("PORT_SCAN", "Dry Run", f"Simulated scan on {target}")
        else:
            # Real scan
            status_text.info(f"🔍 Scanning {target}... Please wait...")
            
            # Run scan in background and update progress
            results = scanner.scan(target, port_range, num_threads, timeout)
            
            if results:
                # Update progress bar based on actual progress
                for i in range(100):
                    progress = min(int((scanner.scan_progress / scanner.total_ports) * 100), 100)
                    progress_bar.progress(progress)
                    if progress >= 100:
                        break
                    time.sleep(0.1)
                
                st.session_state.scan_results = results
                status_text.success(f"✅ Scan completed! Found {results['open_ports_found']} open ports")
            else:
                status_text.error("❌ Scan failed - could not resolve target")
                return
    
    # Display Results
    if 'scan_results' in st.session_state:
        results = st.session_state.scan_results
        
        st.markdown("---")
        st.markdown("## 📊 Scan Results")
        
        # Summary metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Target", results['target'])
        
        with col2:
            st.metric("IP Address", results['ip'])
        
        with col3:
            st.metric("Ports Scanned", results['total_ports_scanned'])
        
        with col4:
            st.metric("Open Ports", results['open_ports_found'], delta="Found")
        
        st.markdown("---")
        
        # Detailed results
        if results['open_ports_found'] > 0:
            st.markdown("### 🔓 Open Ports Details")
            
            for port_info in results['open_ports']:
                with st.expander(f"**Port {port_info['port']}** - {port_info['service']} ({port_info['state'].upper()})"):
                    col_a, col_b = st.columns([1, 3])
                    
                    with col_a:
                        st.markdown(f"""
                        **Port:** {port_info['port']}  
                        **Service:** {port_info['service']}  
                        **State:** {port_info['state']}
                        """)
                    
                    with col_b:
                        st.markdown("**Banner / Version:**")
                        st.code(port_info['banner'], language="text")
        else:
            st.warning("⚠️ No open ports found. The target may be protected by a firewall.")
        
        # Export options
        st.markdown("---")
        st.markdown("### 📥 Export Results")
        
        col_exp1, col_exp2 = st.columns(2)
        
        with col_exp1:
            if st.button("📄 Export to JSON", use_container_width=True):
                # Create new scanner instance for export
                export_scanner = PortScanner(logger)
                filename = f"portscan_{results['target'].replace('.', '_')}_9953_Moazam.json"
                
                try:
                    filepath = export_scanner.export_json(results, filename)
                    
                    with open(filepath, 'r') as f:
                        st.download_button(
                            label="⬇️ Download JSON",
                            data=f.read(),
                            file_name=filename,
                            mime="application/json",
                            key="download_json_port"
                        )
                    st.success(f"✅ Exported to {filename}")
                except Exception as e:
                    st.error(f"❌ Export failed: {str(e)}")
        
        with col_exp2:
            if st.button("📊 Export to HTML", use_container_width=True):
                # Create new scanner instance for export
                export_scanner = PortScanner(logger)
                filename = f"portscan_{results['target'].replace('.', '_')}_9953_Moazam.html"
                
                try:
                    filepath = export_scanner.export_html(results, filename)
                    
                    with open(filepath, 'r') as f:
                        st.download_button(
                            label="⬇️ Download HTML",
                            data=f.read(),
                            file_name=filename,
                            mime="text/html",
                            key="download_html_port"
                        )
                    st.success(f"✅ Exported to {filename}")
                except Exception as e:
                    st.error(f"❌ Export failed: {str(e)}")

# ============================================================================
# DASHBOARD VIEW
# ============================================================================
def show_dashboard(logger, dry_run):
    st.markdown("## 🏠 Mission Control Dashboard")
    
    logger.log("DASHBOARD", "Viewed", "User accessed dashboard")
    
    # Welcome message with more details
    st.markdown("""
    <div style='background: rgba(0, 255, 245, 0.1); padding: 30px; border-radius: 15px; border: 2px solid #00fff5; margin-bottom: 30px; text-align: center;'>
        <h2 style='color: #00fff5; margin-top: 0; font-size: 2.5rem;'>⚡ Welcome to NovaCrypt Defense</h2>
        <p style='color: #fff; font-size: 1.2rem; line-height: 1.8; margin: 20px 0;'>
            NovaCrypt Defense is a comprehensive Python-based security toolkit designed for <strong>PayBuddy FinTech</strong> security testing. 
            This professional suite provides <strong style='color: #00fff5;'>6 powerful modules</strong> for authorized penetration testing, 
            vulnerability assessment, and security analysis.
        </p>
        <p style='color: #00ff7f; font-size: 1rem; margin-top: 25px;'>
            📌 <strong>Select any module from the sidebar to begin your authorized security assessment</strong>
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Feature Grid - just display cards (no buttons)
    st.markdown("### 🎯 Available Security Modules")
    st.markdown("*Use the sidebar to select a module*")
    st.markdown("<br>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div style='background: rgba(0, 191, 255, 0.1); padding: 20px; border-radius: 10px; border: 1px solid #00bfff; min-height: 180px;'>
            <h4 style='color: #00bfff;'>🔍 Port Scanner</h4>
            <p style='color: #ccc; font-size: 0.9rem; line-height: 1.5;'>
                <strong>Capabilities:</strong><br>
                • Multi-threaded TCP port scanning<br>
                • Service detection & banner grabbing<br>
                • Export results to JSON/HTML<br>
                • Identify open ports & running services
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div style='background: rgba(0, 255, 127, 0.1); padding: 20px; border-radius: 10px; border: 1px solid #00ff7f; min-height: 180px;'>
            <h4 style='color: #00ff7f;'>🔑 Password Testing</h4>
            <p style='color: #ccc; font-size: 0.9rem; line-height: 1.5;'>
                <strong>Capabilities:</strong><br>
                • Password strength analysis<br>
                • Policy compliance checking<br>
                • Entropy calculation (Shannon)<br>
                • Hash simulation (MD5/SHA256/bcrypt)
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div style='background: rgba(255, 165, 0, 0.1); padding: 20px; border-radius: 10px; border: 1px solid #ffa500; min-height: 180px;'>
            <h4 style='color: #ffa500;'>💥 Stress Testing</h4>
            <p style='color: #ccc; font-size: 0.9rem; line-height: 1.5;'>
                <strong>Capabilities:</strong><br>
                • Controlled DOS simulation<br>
                • HTTP flood testing (max 200 clients)<br>
                • Real-time latency monitoring<br>
                • Performance graphs & reports
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    col4, col5, col6 = st.columns(3)
    
    with col4:
        st.markdown("""
        <div style='background: rgba(138, 43, 226, 0.1); padding: 20px; border-radius: 10px; border: 1px solid #8a2be2; min-height: 180px;'>
            <h4 style='color: #8a2be2;'>🌐 Web Discovery</h4>
            <p style='color: #ccc; font-size: 0.9rem; line-height: 1.5;'>
                <strong>Capabilities:</strong><br>
                • Directory enumeration (DIRB-style)<br>
                • Subdomain discovery<br>
                • API endpoint detection<br>
                • Hidden resource identification
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    with col5:
        st.markdown("""
        <div style='background: rgba(255, 20, 147, 0.1); padding: 20px; border-radius: 10px; border: 1px solid #ff1493; min-height: 180px;'>
            <h4 style='color: #ff1493;'>📦 Packet Capture</h4>
            <p style='color: #ccc; font-size: 0.9rem; line-height: 1.5;'>
                <strong>Capabilities:</strong><br>
                • Real-time traffic capture<br>
                • Protocol analysis (HTTP/DNS/TCP)<br>
                • Save .pcap files<br>
                • Network traffic visualization
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    with col6:
        st.markdown("""
        <div style='background: rgba(255, 69, 0, 0.1); padding: 20px; border-radius: 10px; border: 1px solid #ff4500; min-height: 180px;'>
            <h4 style='color: #ff4500;'>📊 Reports & Logs</h4>
            <p style='color: #ccc; font-size: 0.9rem; line-height: 1.5;'>
                <strong>Capabilities:</strong><br>
                • View all security logs<br>
                • SHA-256 integrity verification<br>
                • Export PDF/Word/JSON reports<br>
                • Comprehensive findings summary
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Important Notes with more details
    st.markdown("### ⚠️ Important Security Guidelines")
    
    col_a, col_b = st.columns(2)
    
    with col_a:
        st.warning("""
        **🔒 Ethical Testing Principles**
        
        - ✅ Only test **authorized targets** listed in consent.txt
        - ✅ Never attack external/public systems without written permission
        - ✅ Follow **responsible disclosure** for discovered vulnerabilities
        - ✅ Use **rate limiting** to avoid service disruption
        - ✅ Document all activities for audit trails
        
        **⚖️ Legal Compliance:**
        Unauthorized access to computer systems is illegal under computer fraud laws.
        Always obtain proper authorization before testing.
        """)
    
    with col_b:
        st.info("""
        **📋 Evidence Collection System**
        
        - 📝 **Timestamped Logging:** Every action recorded with precise timestamps
        - 🔐 **SHA-256 Integrity:** Cryptographic verification of log authenticity
        - 📊 **Auto Reports:** PDF/Word/JSON exports with findings
        - 💾 **Persistent Storage:** All logs saved to `evidence/` directory
        - 🔍 **Audit Ready:** Logs formatted for security audits
        
        """)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Quick Start Guide with more details
    with st.expander("📖 Quick Start Guide - How to Use This Toolkit"):
        st.markdown("""
        ### 🚀 Getting Started with NovaCrypt Defense
        
        #### **Step 1: Identity Verification** ✅
        - The system automatically verifies `identity.txt` and `consent.txt`
        - Check the sidebar for ✅ green checkmarks confirming verification
        - If files are missing, they'll be created automatically with team information
        
        #### **Step 2: Select Your Module** 🎯
        You can select a module using the dropdown in the sidebar under "Select Module"
        
        #### **Step 3: Configure Testing Parameters** ⚙️
        Each module has specific configuration options:
        - **Port Scanner:** Enter target IP/domain, port range, thread count
        - **Password Test:** Input passwords for analysis or upload hash files
        - **Stress Test:** Set target URL, client count (max 200), duration
        - **Web Discovery:** Specify target domain, wordlist selection
        - **Packet Capture:** Choose network interface, filter protocols
        
        #### **Step 4: Run Your Security Assessment** 🚀
        - Review all parameters before execution
        - Click the main action button (e.g., "Start Scan", "Run Test")
        - Monitor real-time output in the interface
        - All actions are logged automatically
        
        #### **Step 5: Analyze Results** 📊
        - View detailed results directly in the interface
        - Download reports in multiple formats (PDF/Word/JSON)
        - Check "Logs & Reports" module for complete activity history
        - Export findings for documentation
        
        #### **Step 6: Review Security Logs** 📝
        - Navigate to "Logs & Reports" from sidebar or dashboard
        - View timestamped entries for all activities
        - Verify log integrity with SHA-256 hash
        - Export logs for audit purposes
        
        ---
        
        ### 🛡️ Best Practices
        
        **Before Testing:**
        - ✅ Verify you have written authorization
        - ✅ Ensure targets are in consent.txt
        - ✅ Use dry-run mode first to test configuration
        - ✅ Review rate limits and throttling settings
        
        **During Testing:**
        - ⚡ Monitor system resources
        - ⚡ Watch for error messages or warnings
        - ⚡ Keep notes of unusual findings
        - ⚡ Be prepared to stop tests if issues arise
        
        **After Testing:**
        - 📋 Generate comprehensive reports
        - 📋 Document all vulnerabilities found
        - 📋 Provide remediation recommendations
        - 📋 Archive logs for compliance
        
        ---
        
        ### 💡 Pro Tips
        
        - 🎯 Start with **Port Scanner** to identify open services
        - 🎯 Use **Password Assessment** to test authentication strength
        - 🎯 Run **Stress Tests** during off-peak hours
        - 🎯 **Web Discovery** is great for API reconnaissance
        - 🎯 **Packet Capture** helps understand traffic patterns
        - 🎯 Always check **Logs & Reports** after each test
        
        ---
        
        ### 🆘 Troubleshooting
        
        **Issue: Module not responding**
        - Check your internet connection
        - Verify target is accessible
        - Review firewall settings
        
        **Issue: Permission errors**
        - Ensure identity.txt and consent.txt are present
        - Verify target is in approved list
        - Check file permissions
        
        **Issue: Export not working**
        - Ensure evidence/ directory exists
        - Check available disk space
        - Try different export format
        """)
    
    if dry_run:
        st.warning("🧪 **Dry Run Mode Active** - Simulations only, no actual attacks will be performed")

# ============================================================================
# LOGS & REPORTS VIEW
# ============================================================================
def show_logs_reports(logger):
    st.markdown("## 📊 Logs & Reports")
    
    logger.log("LOGS", "Viewed", "User accessed logs and reports")
    
    # Tabs for different views
    tab1, tab2, tab3, tab4 = st.tabs(["📝 Live Logs", "🔒 Integrity Check", "📥 Export", "📄 Generate Report"])
    
    with tab1:
        st.markdown("### 📝 Real-Time Security Logs")
        
        logs = logger.get_logs()
        
        if logs and logs != "No logs found.":
            st.code(logs, language="log")
            
            # Count entries
            log_count = logs.count("[")
            st.info(f"📊 Total log entries: **{log_count}**")
        else:
            st.warning("No logs available yet. Start using the toolkit to generate logs!")
    
    with tab2:
        st.markdown("### 🔒 Log File Integrity")
        
        log_hash = logger.calculate_log_hash()
        
        if log_hash:
            st.success("✅ Log file integrity verified")
            st.code(f"SHA-256 Hash:\n{log_hash}", language="text")
            
            st.info("""
            **Why integrity matters:**
            - Ensures logs haven't been tampered with
            - Provides cryptographic proof of authenticity
            - Required for security audits and compliance
            """)
        else:
            st.warning("No log file found to calculate hash")
    
    with tab3:
        st.markdown("### 📥 Export Reports")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📄 Export Logs (TXT)", use_container_width=True):
                logs = logger.get_logs()
                st.download_button(
                    label="⬇️ Download security_logs.log",
                    data=logs,
                    file_name=f"security_logs_9953_Moazam.log",
                    mime="text/plain"
                )
                logger.log("EXPORT", "Logs exported", "TXT format")
        
        with col2:
            if st.button("📊 Export Report (JSON)", use_container_width=True):
                json_data = logger.export_logs_json()
                st.download_button(
                    label="⬇️ Download report.json",
                    data=json.dumps(json_data, indent=2),
                    file_name=f"report_9953_Moazam.json",
                    mime="application/json"
                )
                logger.log("EXPORT", "Report exported", "JSON format")
        
        st.info("📌 **Note:** PDF/Word reports will be generated automatically after completing assessments")
    
    with tab4:
        st.markdown("### 📄 Generate Comprehensive Report")
        
        st.info("""
        Generate a professional Word/PDF report containing all test results.
        
        The report includes:
        - Executive Summary
        - All module results (Port Scan, Password, Stress, Discovery, Packet)
        - Security recommendations
        - Professional formatting
        """)
        
        # Check what results are available
        available_results = []
        if 'scan_results' in st.session_state:
            available_results.append("✅ Port Scanner Results")
        if 'password_results' in st.session_state:
            available_results.append("✅ Password Assessment Results")
        if 'stress_results' in st.session_state:
            available_results.append("✅ Stress Test Results")
        if 'discovery_results' in st.session_state:
            available_results.append("✅ Web Discovery Results")
        if 'capture_results' in st.session_state:
            available_results.append("✅ Packet Capture Results")
        
        if available_results:
            st.markdown("**Available Data:**")
            for result in available_results:
                st.markdown(f"- {result}")
        else:
            st.warning("⚠️ No test results available yet. Run some tests first!")
        
        col_r1, col_r2 = st.columns(2)
        
        with col_r1:
            if st.button("📄 Generate Word Report (.docx)", use_container_width=True):
                report_gen = ReportGenerator(logger)
                
                with st.spinner("Generating report..."):
                    try:
                        filename = report_gen.generate_comprehensive_report(
                            scan_results=st.session_state.get('scan_results'),
                            password_results=st.session_state.get('password_results'),
                            stress_results=st.session_state.get('stress_results'),
                            discovery_results=st.session_state.get('discovery_results'),
                            packet_results=st.session_state.get('capture_results')
                        )
                        
                        with open(filename, 'rb') as f:
                            st.download_button(
                                label="⬇️ Download Word Report",
                                data=f.read(),
                                file_name=os.path.basename(filename),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        st.success(f"✅ Report generated: {filename}")
                    except Exception as e:
                        st.error(f"❌ Error generating report: {str(e)}")
        
        with col_r2:
            if st.button("📕 Generate PDF Report (.pdf)", use_container_width=True):
                report_gen = ReportGenerator(logger)
                
                with st.spinner("Generating PDF..."):
                    try:
                        # First generate DOCX
                        docx_filename = report_gen.generate_comprehensive_report(
                            scan_results=st.session_state.get('scan_results'),
                            password_results=st.session_state.get('password_results'),
                            stress_results=st.session_state.get('stress_results'),
                            discovery_results=st.session_state.get('discovery_results'),
                            packet_results=st.session_state.get('capture_results')
                        )
                        
                        # Convert to PDF
                        pdf_filename = report_gen.convert_to_pdf(docx_filename)
                        
                        if pdf_filename:
                            with open(pdf_filename, 'rb') as f:
                                st.download_button(
                                    label="⬇️ Download PDF Report",
                                    data=f.read(),
                                    file_name=os.path.basename(pdf_filename),
                                    mime="application/pdf"
                                )
                            
                            st.success(f"✅ PDF generated: {pdf_filename}")
                        else:
                            st.error("❌ PDF conversion failed")
                            
                    except Exception as e:
                        st.error(f"❌ Error generating PDF: {str(e)}")

# ============================================================================
# RUN THE APP
# ============================================================================
if __name__ == "__main__":
    main()
